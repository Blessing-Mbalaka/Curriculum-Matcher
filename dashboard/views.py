from collections import Counter
import csv
from datetime import timedelta
import hashlib
from io import BytesIO
from urllib.parse import urlencode

from django.db import IntegrityError
from django.db.models import Avg
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.urls import reverse
from django.views import View
from django.views.generic import TemplateView, ListView

from courses.models import Course, Module
from jobs.models import JobAdvert
from jobs.ingestion import extract_advert_metadata, extract_advert_sections, parse_advert_date
from analysis.models import AnalysisRun, GapResult, SkillMatrix, TaskRecord
from analysis.spacyskillextraction import classify_skill_text
# Plain functions now — no .delay(), no Celery
from analysis.tasks import (
    run_gap_analysis_task,
    import_csv_task,
    fetch_adzuna_task,
    start_continuous_adzuna_task,
    start_continuous_job_task,
    stop_task,
)


STALE_TASK_MINUTES = 10
BUSINESS_TECHNICAL_SKILL_EXCLUSIONS = {
    "c++",
    "c#",
    "f#",
    "rust",
    "haskell",
    "go",
    "scala",
    "kotlin",
    "java",
    "javascript",
    "typescript",
    "html",
    "css",
    "django",
    "flask",
    "j2ee",
    "z/os",
    "mainframe",
    "kubernetes",
    "devops",
    "ci cd",
    "ci/cd",
    "sre",
    "site reliability engineering",
    "distributed systems",
    "algorithms",
    "encryption",
    "tokenisation",
    "tokenization",
}
BUSINESS_SKILL_LABELS = {
    "data analysis": "business analytics",
    "data analytics": "business analytics",
    "data science": "business analytics",
    "sql": "data-driven decision making",
    "mysql": "data-driven decision making",
    "postgresql": "data-driven decision making",
    "databases": "data governance",
    "python": "analytics automation",
    "r": "analytics automation",
    "machine learning": "AI strategy",
    "artificial intelligence": "AI strategy",
    "generative ai": "AI strategy",
    "genai": "AI strategy",
    "software engineering": "digital transformation",
    "technical architecture": "digital transformation",
    "enterprise architecture": "digital transformation",
    "systems architecture": "digital transformation",
    "cloud native": "digital transformation",
    "cloud-native": "digital transformation",
    "aws": "digital transformation",
    "azure": "digital transformation",
    "gcp": "digital transformation",
    "google cloud": "digital transformation",
    "security": "digital risk governance",
    "web3": "digital transformation",
    "defi": "fintech strategy",
    "fintech": "fintech strategy",
    "automation": "process automation",
    "technical leadership": "digital leadership",
    "technical strategy": "digital strategy",
    "vendor evaluation": "vendor management",
}


def mark_stale_tasks():
    cutoff = timezone.now() - timedelta(minutes=STALE_TASK_MINUTES)
    stale_ids = list(TaskRecord.objects.filter(
        status__in=["PENDING", "STARTED"],
        updated_at__lt=cutoff,
    ).values_list("id", flat=True)[:50])
    if not stale_ids:
        return 0
    return TaskRecord.objects.filter(id__in=stale_ids).update(
        status="FAILURE",
        progress=0,
        notes="Task stopped updating. The background worker likely stopped or the dev server was restarted. Start it again.",
        finished_at=timezone.now(),
    )


def mark_stale_task_if_needed(task):
    if task.status not in ["PENDING", "STARTED"] or not task.updated_at:
        return task
    cutoff = timezone.now() - timedelta(minutes=STALE_TASK_MINUTES)
    if task.updated_at >= cutoff:
        return task
    TaskRecord.objects.filter(
        id=task.id,
        status__in=["PENDING", "STARTED"],
        updated_at__lt=cutoff,
    ).update(
        status="FAILURE",
        progress=0,
        notes="Task stopped updating. The background worker likely stopped or the dev server was restarted. Start it again.",
        finished_at=timezone.now(),
    )
    task.refresh_from_db()
    return task


def task_debug_hint(notes):
    text = (notes or "").lower()
    if "no courses found" in text:
        return "Create a course, add at least one module with content, then start the live pipeline again."
    if "no job adverts found" in text:
        return "Import jobs or let the live pipeline fetch jobs before queueing analysis."
    if "missing adzuna credentials" in text:
        return "Add ADZUNA_APP_ID and ADZUNA_APP_KEY to your .env file, then restart the server."
    if "adzuna api limit reached" in text or "http 429" in text:
        return "The app will retry automatically. Increase the wait interval if this repeats."
    if "adzuna authentication error" in text:
        return "Check ADZUNA_APP_ID and ADZUNA_APP_KEY in .env, then restart the server."
    if "adzuna request error" in text:
        return "Check the keyword, location, and Adzuna country setting."
    if "adzuna network error" in text or "adzuna server error" in text:
        return "This is usually temporary. The jobs-only loop retries automatically."
    if "no text documents" in text:
        return "Add module and job descriptions with enough text for Word2Vec training."
    return "Open Background Tasks for the full task history and check the latest task notes."


def bounded_int(value, default, minimum, maximum):
    try:
        parsed = int(value or default)
    except (TypeError, ValueError):
        parsed = default
    return max(minimum, min(maximum, parsed))


def top_skill_names(counter, limit=5):
    return [skill for skill, _count in counter.most_common(limit)]


def refine_business_skill(skill):
    normalized = (skill or "").strip().lower()
    if normalized in BUSINESS_TECHNICAL_SKILL_EXCLUSIONS:
        return None
    return BUSINESS_SKILL_LABELS.get(normalized, skill)


def refined_business_counter(counter):
    refined = Counter()
    for skill, count in counter.items():
        label = refine_business_skill(skill)
        if label:
            refined[label] += count
    return refined


def refine_skill_rows_for_business(rows, limit=10):
    refined = Counter()
    for row in rows:
        label = refine_business_skill(row["skill"])
        if label:
            refined[label] += row["frequency"]
    return [
        {"skill": skill, "frequency": frequency}
        for skill, frequency in refined.most_common(limit)
    ]


def actual_skill_rows(rows, limit=10):
    return [
        {"skill": row["skill"], "frequency": row["frequency"]}
        for row in rows
        if row.get("skill")
    ][:limit]


def skill_tokens(value):
    stop_words = {"and", "the", "for", "with", "of", "in", "to", "a", "an"}
    clean = "".join(ch.lower() if ch.isalnum() or ch in "+# " else " " for ch in value or "")
    return [token for token in clean.split() if token and token not in stop_words]


def token_cosine_similarity(left, right):
    left_tokens = Counter(skill_tokens(left))
    right_tokens = Counter(skill_tokens(right))
    if not left_tokens or not right_tokens:
        return 0
    dot = sum(left_tokens[token] * right_tokens.get(token, 0) for token in left_tokens)
    left_mag = sum(value * value for value in left_tokens.values()) ** 0.5
    right_mag = sum(value * value for value in right_tokens.values()) ** 0.5
    return dot / max(1e-9, left_mag * right_mag)


MISSING_SECTOR_VALUES = {"", "unclassified", "unknown", "none", "n/a", "na", "other"}
SECTOR_KEYWORDS = [
    ("Cybersecurity and Risk", (
        "cyber", "cybersecurity", "security", "risk", "governance", "compliance",
        "fraud", "audit", "controls", "privacy", "information security",
    )),
    ("Finance and Accounting", (
        "finance", "financial", "accounting", "accountant", "tax", "treasury",
        "investment", "banking", "insurance", "fintech", "payroll", "budget",
    )),
    ("Technology and Data", (
        "data", "analytics", "analysis", "ai", "artificial intelligence",
        "machine learning", "software", "developer", "engineering", "cloud",
        "python", "sql", "database", "digital", "automation", "tableau",
        "power bi", "business intelligence",
    )),
    ("Human Resources and Talent", (
        "human resources", "hr", "talent", "recruitment", "recruiter",
        "learning and development", "training", "people", "workforce",
    )),
    ("Marketing and Sales", (
        "marketing", "sales", "brand", "customer", "crm", "market research",
        "commercial", "business development", "advertising", "media",
    )),
    ("Operations and Supply Chain", (
        "operations", "supply chain", "logistics", "procurement", "warehouse",
        "manufacturing", "production", "process", "quality", "inventory",
    )),
    ("Strategy and Management", (
        "strategy", "strategic", "management", "manager", "leadership",
        "consulting", "project management", "change management", "stakeholder",
        "executive", "mba",
    )),
    ("Education and Curriculum", (
        "course", "curriculum", "module", "teaching", "education", "learning",
        "academic", "university", "programme", "program",
    )),
]


def clean_sector(value):
    sector = " ".join(str(value or "").strip().split())
    return "" if sector.lower() in MISSING_SECTOR_VALUES else sector


def infer_sector_from_text(*values):
    text = " ".join(str(value or "") for value in values).lower()
    if not text.strip():
        return ""
    scores = []
    for sector, keywords in SECTOR_KEYWORDS:
        score = sum(1 for keyword in keywords if keyword in text)
        if score:
            scores.append((score, sector))
    return max(scores, default=(0, ""))[1]


def source_text_for_sector(source_obj, source_label="", parent_label=""):
    values = [source_label, parent_label]
    for field in (
        "title", "description", "summary", "position_info", "category", "company",
        "name", "content",
    ):
        values.append(getattr(source_obj, field, ""))
    course = getattr(source_obj, "course", None)
    if course:
        values.extend([getattr(course, "code", ""), getattr(course, "name", ""), getattr(course, "description", "")])
    return " ".join(str(value or "") for value in values)


def join_skill_names(skills):
    if not skills:
        return ""
    if len(skills) == 1:
        return skills[0]
    return f"{', '.join(skills[:-1])} and {skills[-1]}"


def skill_entity_id(skill):
    canonical = " ".join((skill or "").lower().replace("-", " ").split())
    slug = "".join(ch if ch.isalnum() else "-" for ch in canonical).strip("-")
    slug = "-".join(part for part in slug.split("-") if part)
    digest = hashlib.sha1(canonical.encode("utf-8")).hexdigest()[:8]
    normalized_slug_source = "-".join(canonical.split())
    if slug and slug != normalized_slug_source:
        return f"skill-{slug[:63]}-{digest}"
    if slug:
        return f"skill-{slug[:72]}"
    return f"skill-{digest}"


def normalize_skill_entity(entity_or_skill):
    if isinstance(entity_or_skill, dict):
        skill = entity_or_skill.get("skill") or entity_or_skill.get("text") or ""
        entity_id = entity_or_skill.get("id") or skill_entity_id(skill)
        return {
            "id": entity_id,
            "skill": skill,
            "source": entity_or_skill.get("source") or "ner",
            "confidence": entity_or_skill.get("confidence"),
            "mention_count": entity_or_skill.get("mention_count") or len(entity_or_skill.get("mentions") or []) or 1,
        }
    skill = str(entity_or_skill or "").strip()
    return {
        "id": skill_entity_id(skill),
        "skill": skill,
        "source": "legacy",
        "confidence": None,
        "mention_count": 1,
    }


def job_skill_entities(job, fallback_skills=None, limit=10):
    raw_entities = getattr(job, "skill_entities", None) or []
    if not raw_entities:
        raw_entities = getattr(job, "skills_extracted", None) or fallback_skills or []
    seen = set()
    entities = []
    for raw in raw_entities:
        entity = normalize_skill_entity(raw)
        if not entity["skill"] or entity["id"] in seen:
            continue
        seen.add(entity["id"])
        entities.append(entity)
        if len(entities) >= limit:
            break
    return entities


def entity_row(source_type, source_obj, entity, source_label, parent_label=""):
    normalized = normalize_skill_entity(entity)
    if isinstance(entity, dict):
        normalized.update({
            "chunk_id": entity.get("chunk_id") or f"{source_type}-{source_obj.id}-{normalized['id']}",
            "label": entity.get("label") or "SKILL",
            "tier": entity.get("tier") or "explicit",
            "skill_type": entity.get("skill_type") or "domain",
            "pattern": entity.get("pattern") or "",
            "pos_signature": entity.get("pos_signature") or "",
            "text": entity.get("text") or normalized["skill"],
            "start": entity.get("start"),
            "end": entity.get("end"),
            "label_status": entity.get("label_status") or "machine",
        })
    else:
        classification = classify_skill_text(normalized["skill"])
        normalized.update({
            "chunk_id": f"{source_type}-{source_obj.id}-{normalized['id']}",
            "label": "SKILL",
            "tier": classification["tier"],
            "skill_type": classification["skill_type"],
            "classification_scores": classification["scores"],
            "pattern": "",
            "pos_signature": "",
            "text": normalized["skill"],
            "start": None,
            "end": None,
            "label_status": "legacy",
        })
    extracted_date = getattr(source_obj, "created_at", None)
    if isinstance(source_obj, JobAdvert):
        extracted_date = source_obj.date_posted or source_obj.created_at
    explicit_sector = ""
    if isinstance(entity, dict):
        label = (entity.get("label") or "").lower()
        if label in {"sector", "industry"}:
            explicit_sector = entity.get("skill") or entity.get("text") or ""
    source_sector_text = source_text_for_sector(source_obj, source_label, parent_label)
    inferred_sector = infer_sector_from_text(source_sector_text, normalized.get("skill"), normalized.get("text"))
    if isinstance(source_obj, JobAdvert):
        sector = clean_sector(explicit_sector) or clean_sector(source_obj.category) or inferred_sector
        job_title = source_obj.title
    else:
        job_title = ""
        sector = clean_sector(explicit_sector) or inferred_sector or clean_sector(parent_label) or "Course module"
    sector = clean_sector(sector) or "Unclassified"

    normalized.update({
        "source_type": source_type,
        "source_id": source_obj.id,
        "source_label": source_label,
        "parent_label": parent_label,
        "sector": sector,
        "job_title": job_title,
        "extracted_date": extracted_date.isoformat() if extracted_date else "",
        "extracted_year": extracted_date.year if extracted_date else None,
        "unique_key": f"{source_type}:{source_obj.id}:{normalized['id']}",
    })
    return normalized


def iter_skill_entity_rows():
    for job in JobAdvert.objects.order_by("title", "id"):
        raw_entities = job.skill_entities or job.skills_extracted or []
        source_label = f"{job.title} @ {job.company}" if job.company else job.title
        for entity in raw_entities:
            yield entity_row("job", job, entity, source_label)

    modules = Module.objects.select_related("course").order_by("course__code", "order", "name", "id")
    for module in modules:
        raw_entities = module.skill_entities or module.skills_extracted or []
        source_label = module.name
        parent_label = module.course.code or module.course.name
        for entity in raw_entities:
            yield entity_row("module", module, entity, source_label, parent_label)


def filtered_skill_entity_rows(request):
    rows = list(iter_skill_entity_rows())
    source = request.GET.get("source", "").strip()
    skill_type = request.GET.get("skill_type", "").strip()
    sector = request.GET.get("sector", "").strip()
    job_title = request.GET.get("job_title", "").strip()
    q = request.GET.get("q", "").strip().lower()
    if source:
        rows = [row for row in rows if row["source_type"] == source]
    if skill_type:
        rows = [row for row in rows if row["skill_type"] == skill_type]
    if sector:
        rows = [row for row in rows if row["sector"] == sector]
    if job_title:
        rows = [row for row in rows if row["job_title"] == job_title]
    if q:
        rows = [
            row for row in rows
            if q in row["skill"].lower()
            or q in row["source_label"].lower()
            or q in row["sector"].lower()
            or q in row["job_title"].lower()
            or q in row["chunk_id"].lower()
            or q in row["id"].lower()
        ]
    return rows


def update_entity_collection(obj, entity_id, chunk_id, updates):
    entities = list(getattr(obj, "skill_entities", None) or [])
    if not entities and getattr(obj, "skills_extracted", None):
        source_type = "job" if isinstance(obj, JobAdvert) else "module"
        entities = [
            # Preserve stable legacy chunk IDs while upgrading the row to editable entity data.
            {
                "id": skill_entity_id(skill),
                "chunk_id": f"{source_type}-{obj.id}-{skill_entity_id(skill)}",
                "skill": skill,
                "label": "SKILL",
                "tier": classify_skill_text(skill)["tier"],
                "skill_type": classify_skill_text(skill)["skill_type"],
                "classification_scores": classify_skill_text(skill)["scores"],
                "source": "legacy",
                "confidence": None,
                "mention_count": 1,
                "label_status": "legacy",
            }
            for skill in obj.skills_extracted
        ]
    changed = False
    for entity in entities:
        if not isinstance(entity, dict):
            continue
        if entity.get("chunk_id") == chunk_id or entity.get("id") == entity_id:
            entity.update({key: value for key, value in updates.items() if value != ""})
            entity["label_status"] = "reviewed"
            changed = True
            break
    if not changed:
        return False
    obj.skill_entities = entities
    obj.skills_extracted = sorted({entity.get("skill") for entity in entities if isinstance(entity, dict) and entity.get("skill")})
    obj.save(update_fields=["skill_entities", "skills_extracted"])
    return True


def recommendation_skill_insight(missing_counter, matched_counter):
    missing = top_skill_names(refined_business_counter(missing_counter))
    matched = top_skill_names(refined_business_counter(matched_counter), 4)
    parts = []
    if missing:
        parts.append(f"Based on the analysed course-to-job data, the curriculum should strengthen {join_skill_names(missing)}.")
    if matched:
        parts.append(f"The current evidence already shows coverage in {join_skill_names(matched)}.")
    if not parts:
        parts.append("No repeated skill pattern is visible yet, so inspect the lowest-scoring job matches first.")
    return " ".join(parts)


def build_skill_suggestion_matrix(results, limit=8):
    matrix = {}
    for result in results:
        matched_labels = {
            label for label in (refine_business_skill(skill) for skill in (result.matched_skills or []))
            if label
        }
        missing_labels = {
            label for label in (refine_business_skill(skill) for skill in (result.missing_skills or []))
            if label
        }
        for skill in matched_labels | missing_labels:
            item = matrix.setdefault(skill, {
                "job_ids": set(),
                "covered_course_ids": set(),
                "matched_pairs": set(),
                "missing_pairs": set(),
            })
            item["job_ids"].add(result.job_id)
            pair = (result.course_id, result.job_id)
            if skill in matched_labels:
                item["covered_course_ids"].add(result.course_id)
                item["matched_pairs"].add(pair)
            if skill in missing_labels:
                item["missing_pairs"].add(pair)

    max_gap = max((len(item["missing_pairs"]) for item in matrix.values()), default=1)
    rows = []
    for skill, item in matrix.items():
        matched_count = len(item["matched_pairs"])
        missing_count = len(item["missing_pairs"])
        demand_count = matched_count + missing_count
        coverage = round((matched_count / max(1, demand_count)) * 100)
        gap_percent = round((missing_count / max(1, demand_count)) * 100)
        rows.append({
            "skill": skill,
            "demand_count": demand_count,
            "job_count": len(item["job_ids"]),
            "course_count": len(item["covered_course_ids"]),
            "matched_count": matched_count,
            "missing_count": missing_count,
            "coverage": coverage,
            "gap_percent": gap_percent,
            "gap_width": round((missing_count / max_gap) * 100),
        })
    return sorted(
        rows,
        key=lambda row: (-row["missing_count"], row["coverage"], row["skill"]),
    )[:limit]


class DashboardView(TemplateView):
    template_name = "dashboard/home.html"

    def get_context_data(self, **kwargs):
        mark_stale_tasks()
        ctx = super().get_context_data(**kwargs)
        ctx["course_count"] = Course.objects.count()
        ctx["module_count"] = sum(c.modules.count() for c in Course.objects.prefetch_related("modules"))
        ctx["job_count"] = JobAdvert.objects.count()
        ctx["last_run"] = AnalysisRun.objects.first()
        ctx["pending_tasks"] = TaskRecord.objects.filter(status__in=["PENDING", "STARTED"]).count()
        ctx["live_task"] = (
            TaskRecord.objects
            .filter(run_name__startswith="Jobs Only", status__in=["PENDING", "STARTED"])
            .first()
            or TaskRecord.objects
            .filter(run_name__startswith="Live Pipeline", status__in=["PENDING", "STARTED"])
            .first()
        )
        latest_jobs_only = TaskRecord.objects.filter(run_name__startswith="Jobs Only").order_by("-created_at").first()
        ctx["should_autostart_jobs"] = not ctx["live_task"] and not (latest_jobs_only and latest_jobs_only.status == "STOPPED")
        ctx["avg_score"] = (GapResult.objects.aggregate(v=Avg("similarity_score"))["v"] or 0) * 100
        ctx["latest_results"] = GapResult.objects.select_related("course", "job").order_by("-run__created_at", "-similarity_score")[:8]
        ctx["network_schools"] = (
            Course.objects
            .exclude(university_name="")
            .order_by("university_name")
            .values_list("university_name", flat=True)
            .distinct()
        )
        ctx["network_jobs"] = JobAdvert.objects.order_by("title").values("id", "title", "company")[:500]
        return ctx


class DataExportView(TemplateView):
    template_name = "dashboard/data_export.html"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        all_rows = list(iter_skill_entity_rows())
        rows = filtered_skill_entity_rows(self.request)
        type_counts = Counter(row["skill_type"] for row in rows)
        source_counts = Counter(row["source_type"] for row in rows)
        all_source_counts = Counter(row["source_type"] for row in all_rows)
        current_source = self.request.GET.get("source", "")
        current_skill_type = self.request.GET.get("skill_type", "")
        current_sector = self.request.GET.get("sector", "")
        current_job_title = self.request.GET.get("job_title", "")
        current_q = self.request.GET.get("q", "")
        tile_query = {
            "skill_type": current_skill_type,
            "q": current_q,
        }
        tile_query = {key: value for key, value in tile_query.items() if value}
        job_tile_query = {**tile_query, "source": "job"}
        module_tile_query = {**tile_query, "source": "module"}
        export_query = {
            "source": current_source,
            "skill_type": current_skill_type,
            "sector": current_sector,
            "job_title": current_job_title,
            "q": current_q,
        }
        export_query = {key: value for key, value in export_query.items() if value}
        ctx.update({
            "rows": rows[:500],
            "total_rows": len(rows),
            "type_counts": type_counts,
            "source_counts": source_counts,
            "all_source_counts": all_source_counts,
            "all_skill_rows_count": len(all_rows),
            "current_source": current_source,
            "current_skill_type": current_skill_type,
            "current_sector": current_sector,
            "current_job_title": current_job_title,
            "current_q": current_q,
            "skill_export_href": f"{reverse('skill-entity-export')}?{urlencode(export_query)}" if export_query else reverse("skill-entity-export"),
            "visual_export_href": f"{reverse('data-export-visual-export')}?{urlencode(export_query)}" if export_query else reverse("data-export-visual-export"),
            "skill_vector_href": f"{reverse('skill-vector-space')}?{urlencode(export_query)}" if export_query else reverse("skill-vector-space"),
            "source_tiles": [
                {
                    "key": "job",
                    "label": "Job Skills",
                    "count": all_source_counts.get("job", 0),
                    "active": current_source == "job",
                    "href": f"{reverse('data-export')}?{urlencode(job_tile_query)}",
                },
                {
                    "key": "module",
                    "label": "Course Skills",
                    "count": all_source_counts.get("module", 0),
                    "active": current_source == "module",
                    "href": f"{reverse('data-export')}?{urlencode(module_tile_query)}",
                },
                {
                    "key": "merged",
                    "label": "Merged Skills",
                    "count": len(all_rows),
                    "active": not current_source,
                    "href": f"{reverse('data-export')}?{urlencode(tile_query)}" if tile_query else reverse("data-export"),
                },
            ],
            "skill_types": sorted({row["skill_type"] for row in all_rows if row["skill_type"]}),
            "sectors": sorted({row["sector"] for row in all_rows if row["sector"]}),
            "job_titles": sorted({row["job_title"] for row in all_rows if row["job_title"]}),
        })
        return ctx


class SkillVectorSpaceView(TemplateView):
    template_name = "dashboard/skill_vector_space.html"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        query_string = self.request.META.get("QUERY_STRING", "")
        api_url = reverse("skill-vector-space-api")
        csv_url = reverse("skill-vector-space-export")
        if query_string:
            api_url = f"{api_url}?{query_string}"
            csv_url = f"{csv_url}?{query_string}"
        ctx.update({
            "api_url": api_url,
            "csv_url": csv_url,
            "back_url": f"{reverse('data-export')}?{query_string}" if query_string else reverse("data-export"),
        })
        return ctx


class SkillEntityUpdateView(View):
    def post(self, request):
        source_type = request.POST.get("source_type", "")
        source_id = bounded_int(request.POST.get("source_id"), 0, 0, 100000000)
        entity_id = request.POST.get("entity_id", "")
        chunk_id = request.POST.get("chunk_id", "")
        updates = {
            "skill": request.POST.get("skill", "").strip(),
            "label": request.POST.get("label", "").strip() or "SKILL",
            "tier": request.POST.get("tier", "").strip() or "reviewed",
            "skill_type": request.POST.get("skill_type", "").strip() or "domain",
        }
        model = JobAdvert if source_type == "job" else Module if source_type == "module" else None
        if not model or not source_id:
            messages.error(request, "Could not update skill label because the source row was invalid.")
            return redirect("data-export")
        obj = get_object_or_404(model, pk=source_id)
        if update_entity_collection(obj, entity_id, chunk_id, updates):
            messages.success(request, "Skill label updated.")
        else:
            messages.error(request, "Could not find that skill chunk to update.")
        next_url = request.POST.get("next") or reverse("data-export")
        return redirect(next_url)


class SkillEntityCsvExportView(View):
    def get(self, request):
        rows = filtered_skill_entity_rows(request)
        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = 'attachment; filename="skill-entities.csv"'
        writer = csv.writer(response)
        writer.writerow([
            "chunk_id", "entity_id", "skill", "label", "tier", "skill_type",
            "source_type", "source_id", "source_label", "parent_label",
            "sector", "job_title", "extracted_date", "extracted_year",
            "source", "confidence", "mention_count", "pattern", "pos_signature",
            "text", "start", "end", "label_status", "unique_key",
        ])
        for row in rows:
            writer.writerow([
                row["chunk_id"], row["id"], row["skill"], row["label"], row["tier"], row["skill_type"],
                row["source_type"], row["source_id"], row["source_label"], row["parent_label"],
                row["sector"], row["job_title"], row["extracted_date"], row["extracted_year"],
                row["source"], row["confidence"], row["mention_count"], row["pattern"], row["pos_signature"],
                row["text"], row["start"], row["end"], row["label_status"], row["unique_key"],
            ])
        return response


class JobCsvExportView(View):
    def get(self, request):
        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = 'attachment; filename="jobs-with-skills.csv"'
        writer = csv.writer(response)
        writer.writerow([
            "job_id", "title", "company", "location", "category", "source",
            "date_posted", "skill_count", "skills", "skill_entity_ids",
        ])
        for job in JobAdvert.objects.order_by("title", "id"):
            entities = [normalize_skill_entity(entity) for entity in (job.skill_entities or job.skills_extracted or [])]
            writer.writerow([
                job.id,
                job.title,
                job.company,
                job.location,
                job.category,
                job.source,
                job.date_posted.isoformat() if job.date_posted else "",
                len({entity["id"] for entity in entities}),
                "; ".join(sorted({entity["skill"] for entity in entities if entity["skill"]})),
                "; ".join(sorted({entity["id"] for entity in entities if entity["id"]})),
            ])
        return response


def csv_download_response(filename):
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


class DataExportVisualCsvExportView(View):
    def get(self, request):
        rows = filtered_skill_entity_rows(request)
        response = csv_download_response("data-export-visual-source.csv")
        writer = csv.writer(response)
        writer.writerow(["visual", "dimension_1", "dimension_2", "dimension_3", "value", "notes"])

        for skill, count in Counter(row["skill"] for row in rows).most_common():
            writer.writerow(["top_skill_evidence", skill, "", "", count, "Count of filtered skill evidence records"])
        for skill_type, count in Counter(row["skill_type"] for row in rows).most_common():
            writer.writerow(["skill_type_donut", skill_type, "", "", count, "Count of filtered records by skill type"])
        for source_type, count in Counter(row["source_type"] for row in rows).most_common():
            writer.writerow(["source_bar", source_type, "", "", count, "Count of filtered records by source"])
        type_source_counts = Counter((row["source_type"], row["skill_type"]) for row in rows)
        for (source_type, skill_type), count in sorted(type_source_counts.items()):
            writer.writerow(["skill_heatmap", source_type, skill_type, "", count, "Source by skill-type count"])
        forecast_counts = Counter((row["skill"], row["extracted_year"]) for row in rows if row.get("extracted_year"))
        for (skill, year), count in sorted(forecast_counts.items(), key=lambda item: (item[0][0], item[0][1])):
            writer.writerow(["skill_forecast", skill, year, "", count, "Yearly count used by forecast visual"])

        top_skill_rows = Counter(row["skill"] for row in rows).most_common(36)
        top_skills = [skill for skill, _count in top_skill_rows]
        skill_context = {
            skill: {
                "rows": [row for row in rows if row["skill"] == skill],
                "tokens": skill_tokens(skill),
            }
            for skill in top_skills
        }
        for left_index, left_skill in enumerate(top_skills):
            for right_skill in top_skills[left_index + 1:]:
                left_rows = skill_context[left_skill]["rows"]
                right_rows = skill_context[right_skill]["rows"]
                token_score = token_cosine_similarity(left_skill, right_skill)
                type_score = 0.18 if Counter(row["skill_type"] for row in left_rows).most_common(1)[0][0] == Counter(row["skill_type"] for row in right_rows).most_common(1)[0][0] else 0
                left_sector = Counter(row["sector"] for row in left_rows).most_common(1)[0][0]
                right_sector = Counter(row["sector"] for row in right_rows).most_common(1)[0][0]
                sector_score = 0.12 if left_sector == right_sector else 0
                score = round(token_score + type_score + sector_score, 3)
                if score >= 0.28:
                    writer.writerow(["semantic_skill_cluster", left_skill, right_skill, "association_score", score, "Token, skill-type, and sector association"])
        return response


class JobUploadView(View):
    template_name = "jobs/upload.html"

    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        t = request.POST.get("upload_type", "csv")

        if t == "csv":
            f = request.FILES.get("csv_file")
            if not f:
                messages.error(request, "Select a CSV file first.")
                return render(request, self.template_name)
            raw = f.read()          # bytes — passed directly, no list() conversion needed
            record = TaskRecord.objects.create(run_name=f"CSV: {f.name}")
            import_csv_task(raw, record_id=record.id)   # fires thread, returns immediately
            messages.success(request, f"'{f.name}' is being imported in the background — track progress under Tasks.")
            return redirect("task-list")

        elif t == "adzuna":
            kw = request.POST.get("keyword", "").strip()
            loc = request.POST.get("location", "south africa").strip()
            n = int(request.POST.get("max_results", 50))
            if not kw:
                messages.error(request, "Enter a keyword.")
                return render(request, self.template_name)
            record = TaskRecord.objects.create(run_name=f"Adzuna: {kw}")
            fetch_adzuna_task(kw, loc, n, record_id=record.id)
            messages.success(request, f"Fetching '{kw}' jobs in the background.")
            return redirect("task-list")

        elif t == "manual":
            title = request.POST.get("title", "").strip()
            desc = request.POST.get("description", "").strip()
            if not title or not desc:
                messages.error(request, "Title and description are required.")
                return render(request, self.template_name)
            sections = extract_advert_sections(desc)
            metadata = extract_advert_metadata(desc)
            try:
                JobAdvert.objects.create(
                    title=title,
                    company=request.POST.get("company", "").strip(),
                    recruiter=request.POST.get("recruiter", "").strip() or metadata.get("recruiter", ""),
                    job_reference=request.POST.get("job_reference", "").strip() or metadata.get("job_reference", ""),
                    location=request.POST.get("location", "").strip() or metadata.get("location", ""),
                    category=request.POST.get("category", "").strip(),
                    contract_type=request.POST.get("contract_type", "").strip(),
                    contract_time=request.POST.get("contract_time", "").strip(),
                    summary=request.POST.get("summary", "").strip() or sections["summary"],
                    position_info=request.POST.get("position_info", "").strip() or sections["position_info"],
                    description=desc,
                    source="upload",
                    date_posted=parse_advert_date(metadata.get("date_posted")),
                )
            except IntegrityError:
                messages.warning(request, "That job advert already exists, so it was not added again.")
                return redirect("job-list")
            messages.success(request, f"Job '{title}' added.")
            return redirect("job-list")

        return redirect("job-list")


class JobListView(ListView):
    model = JobAdvert
    template_name = "jobs/list.html"
    context_object_name = "jobs"
    paginate_by = 30


class JobDeleteView(View):
    def post(self, request, pk):
        get_object_or_404(JobAdvert, pk=pk).delete()
        messages.success(request, "Job deleted.")
        return redirect("job-list")


class RunAnalysisView(View):
    def post(self, request):
        name = request.POST.get("run_name", "Analysis Run")
        max_jobs = bounded_int(request.POST.get("max_jobs"), 0, 0, 100000) or None
        if max_jobs and "Smoke" not in name:
            name = f"Smoke {name}"
        record = TaskRecord.objects.create(run_name=name)
        run_gap_analysis_task(run_name=name, record_id=record.id, max_jobs=max_jobs)   # fires thread
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({
                "task_id": record.id,
                "status_url": reverse("task-status-api", args=[record.id]),
                "task_url": reverse("task-list"),
                "max_jobs": max_jobs,
            })
        messages.success(request, f"Analysis '{name}' started in the background.")
        return redirect("task-list")


class StartContinuousJobsView(View):
    def post(self, request):
        mark_stale_tasks()
        live = TaskRecord.objects.filter(run_name__startswith="Live Pipeline", status__in=["PENDING", "STARTED"]).first()
        if live:
            return JsonResponse({
                "task_id": live.id,
                "status_url": reverse("task-status-api", args=[live.id]),
            })

        keyword = request.POST.get("keyword", "MBA").strip() or "MBA"
        location = request.POST.get("location", "south africa").strip() or "south africa"
        max_results = bounded_int(request.POST.get("max_results"), 50, 1, 50)
        interval = bounded_int(request.POST.get("interval_seconds"), 45, 10, 3600)
        record = TaskRecord.objects.create(run_name=f"Live Pipeline: {keyword}")
        start_continuous_job_task(keyword, location, max_results, interval, record_id=record.id)
        return JsonResponse({
            "task_id": record.id,
            "status_url": reverse("task-status-api", args=[record.id]),
        })


class StartJobsOnlyView(View):
    def post(self, request):
        mark_stale_tasks()
        live = TaskRecord.objects.filter(run_name__startswith="Jobs Only", status__in=["PENDING", "STARTED"]).first()
        if live:
            return JsonResponse({
                "task_id": live.id,
                "status_url": reverse("task-status-api", args=[live.id]),
            })

        keyword = request.POST.get("keyword", "MBA").strip() or "MBA"
        location = request.POST.get("location", "south africa").strip() or "south africa"
        max_results = bounded_int(request.POST.get("max_results"), 50, 1, 50)
        interval = bounded_int(request.POST.get("interval_seconds"), 45, 5, 3600)
        record = TaskRecord.objects.create(run_name=f"Jobs Only: {keyword}")
        start_continuous_adzuna_task(keyword, location, max_results, interval, record_id=record.id)
        return JsonResponse({
            "task_id": record.id,
            "status_url": reverse("task-status-api", args=[record.id]),
        })


class StopTaskView(View):
    def post(self, request, pk):
        stop_task(pk)
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": True})
        messages.success(request, "Pause requested.")
        return redirect("home")


class AnalysisResultsView(TemplateView):
    template_name = "analysis/results.html"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        runs = AnalysisRun.objects.all()
        ctx["runs"] = runs

        run_id = self.request.GET.get("run")
        sel = None
        if run_id:
            sel = get_object_or_404(AnalysisRun, id=run_id)
        elif runs.exists():
            sel = runs.first()

        if sel:
            school_filter = self.request.GET.get("school", "").strip()
            threshold = bounded_int(self.request.GET.get("threshold"), 55, 0, 100)
            all_results_qs = (
                GapResult.objects
                .filter(run=sel)
                .select_related("course", "job")
                .order_by("-similarity_score")
            )
            if school_filter:
                all_results_qs = all_results_qs.filter(course__university_name=school_filter)
            all_results = list(all_results_qs)

            ctx["selected_run"] = sel
            ctx["selected_school"] = school_filter
            ctx["threshold"] = threshold
            ctx["schools"] = (
                Course.objects
                .filter(gapresult__run=sel)
                .exclude(university_name="")
                .order_by("university_name")
                .values_list("university_name", flat=True)
                .distinct()
            )
            ctx["results"] = all_results[:300]
            ctx["job_skills"] = SkillMatrix.objects.filter(run=sel, source="jobs")[:20]
            ctx["course_skills"] = SkillMatrix.objects.filter(run=sel, source="courses")[:20]
            visual_data = build_results_visual_data(all_results, threshold)
            ctx.update(visual_data)
        return ctx


def build_results_visual_data(results, threshold):
    bands = [
        {"key": "0-20", "low": 0, "high": 20},
        {"key": "20-40", "low": 20, "high": 40},
        {"key": "40-60", "low": 40, "high": 60},
        {"key": "60-80", "low": 60, "high": 80},
        {"key": "80-100", "low": 80, "high": 101},
    ]
    course_map = {}
    school_map = {}
    max_cell = 1

    for result in results:
        course = result.course
        school = course.university_name or "Unassigned school"
        score = result.similarity_percent
        course_item = course_map.setdefault(course.id, {
            "course": course,
            "school": school,
            "scores": [],
            "matched_total": 0,
            "missing_total": 0,
            "matched_skills": Counter(),
            "missing_skills": Counter(),
            "cells": {band["key"]: 0 for band in bands},
        })
        course_item["scores"].append(score)
        course_item["matched_total"] += len(result.matched_skills or [])
        course_item["missing_total"] += len(result.missing_skills or [])
        course_item["matched_skills"].update(result.matched_skills or [])
        course_item["missing_skills"].update(result.missing_skills or [])
        for band in bands:
            if band["low"] <= score < band["high"]:
                course_item["cells"][band["key"]] += 1
                max_cell = max(max_cell, course_item["cells"][band["key"]])
                break

        school_item = school_map.setdefault(school, {
            "scores": [],
            "matched_total": 0,
            "missing_total": 0,
            "matched_skills": Counter(),
            "missing_skills": Counter(),
            "courses": set(),
        })
        school_item["scores"].append(score)
        school_item["matched_total"] += len(result.matched_skills or [])
        school_item["missing_total"] += len(result.missing_skills or [])
        school_item["matched_skills"].update(result.matched_skills or [])
        school_item["missing_skills"].update(result.missing_skills or [])
        school_item["courses"].add(course.id)

    heatmap_rows = []
    plotly_heatmap_rows = []
    scatter_points = []
    course_recommendations = []
    for item in sorted(course_map.values(), key=lambda value: value["course"].name):
        scores = item["scores"]
        avg_score = round(sum(scores) / max(1, len(scores)), 1)
        mismatch = avg_score < threshold
        cells = []
        for band in bands:
            count = item["cells"][band["key"]]
            intensity = count / max_cell if max_cell else 0
            cells.append({"band": band["key"], "count": count, "intensity": round(intensity, 3)})
        heatmap_rows.append({
            "course": item["course"],
            "school": item["school"],
            "avg_score": avg_score,
            "mismatch": mismatch,
            "cells": cells,
        })
        plotly_heatmap_rows.extend([
            {
                "course_id": result.course_id,
                "course": result.course.code or result.course.name,
                "course_name": result.course.name,
                "school": result.course.university_name or "Unassigned school",
                "job_id": result.job_id,
                "job": result.job.title,
                "company": result.job.company or "",
                "score": result.similarity_percent,
                "matched_count": len(result.matched_skills or []),
                "missing_count": len(result.missing_skills or []),
                "matched": ", ".join((result.matched_skills or [])[:8]),
                "missing": ", ".join((result.missing_skills or [])[:8]),
            }
            for result in results
            if result.course_id == item["course"].id
        ])
        scatter_points.append({
            "label": item["course"].code or item["course"].name,
            "course": item["course"].name,
            "school": item["school"],
            "x": item["matched_total"],
            "y": item["missing_total"],
            "score": avg_score,
        })
        if mismatch:
            skill_insight = recommendation_skill_insight(item["missing_skills"], item["matched_skills"])
            suggested_skills = join_skill_names(top_skill_names(refined_business_counter(item["missing_skills"])))
            course_recommendations.append({
                "label": item["course"].code or item["course"].name,
                "school": item["school"],
                "score": avg_score,
                "skills": suggested_skills,
                "message": f"For this programme, the data suggests a curriculum refresh around {suggested_skills or 'the lowest-covered demand skills'}. {skill_insight} Review module outcomes, readings, projects, and assessment language around those demand areas.",
            })

    school_recommendations = []
    school_summaries = []
    for school, item in sorted(school_map.items()):
        scores = item["scores"]
        avg_score = round(sum(scores) / max(1, len(scores)), 1)
        matched = item["matched_total"]
        missing = item["missing_total"]
        mismatch = avg_score < threshold or missing > matched
        summary = {
            "school": school,
            "avg_score": avg_score,
            "matched_total": matched,
            "missing_total": missing,
            "course_count": len(item["courses"]),
            "mismatch": mismatch,
        }
        school_summaries.append(summary)
        if mismatch:
            if avg_score < threshold:
                reason = f"Average score is {avg_score}%, below the {threshold}% threshold."
            else:
                reason = f"Missing skill evidence ({missing}) exceeds matched evidence ({matched})."
            skill_insight = recommendation_skill_insight(item["missing_skills"], item["matched_skills"])
            school_recommendations.append({
                "school": school,
                "score": avg_score,
                "message": f"{reason} {skill_insight} Prioritise curriculum updates in modules linked to the job adverts with the highest missing-skill counts.",
            })

    if not school_recommendations and school_summaries:
        best = max(school_summaries, key=lambda item: item["avg_score"])
        school_recommendations.append({
            "school": best["school"],
            "score": best["avg_score"],
            "message": "No curriculum is below the current threshold. Use the scatter plot to inspect outlier courses with high missing-skill counts before changing module content.",
        })

    return {
        "score_bands": [band["key"] for band in bands],
        "heatmap_rows": heatmap_rows,
        "plotly_heatmap_rows": plotly_heatmap_rows[:500],
        "scatter_points": scatter_points,
        "school_summaries": school_summaries,
        "school_recommendations": school_recommendations,
        "course_recommendations": course_recommendations[:12],
        "skill_suggestion_matrix": build_skill_suggestion_matrix(results),
    }


def filtered_analysis_results_from_request(request):
    run_id = request.GET.get("run")
    runs = AnalysisRun.objects.all()
    selected_run = get_object_or_404(AnalysisRun, id=run_id) if run_id else runs.first()
    if not selected_run:
        return None, [], bounded_int(request.GET.get("threshold"), 55, 0, 100)
    threshold = bounded_int(request.GET.get("threshold"), 55, 0, 100)
    qs = (
        GapResult.objects
        .filter(run=selected_run)
        .select_related("course", "job")
        .order_by("-similarity_score")
    )
    school_filter = request.GET.get("school", "").strip()
    if school_filter:
        qs = qs.filter(course__university_name=school_filter)
    return selected_run, list(qs), threshold


class AnalysisVisualCsvExportView(View):
    def get(self, request):
        selected_run, results, threshold = filtered_analysis_results_from_request(request)
        response = csv_download_response("gap-analysis-visual-source.csv")
        writer = csv.writer(response)
        writer.writerow([
            "visual", "run", "course_id", "course", "school", "job_id", "job", "company",
            "dimension", "value", "score_percent", "matched_count", "missing_count",
            "matched_skills", "missing_skills", "notes",
        ])
        run_name = selected_run.name if selected_run else ""
        visual_data = build_results_visual_data(results, threshold)

        for row in visual_data["plotly_heatmap_rows"]:
            writer.writerow([
                "course_to_job_gap_heatmap", run_name, row["course_id"], row["course_name"], row["school"],
                row["job_id"], row["job"], row["company"], "match_score_percent", row["score"], row["score"],
                row["matched_count"], row["missing_count"], row["matched"], row["missing"],
                "One course-to-job cell used by ranked and matrix Plotly views",
            ])
        for row in visual_data["heatmap_rows"]:
            for cell in row["cells"]:
                writer.writerow([
                    "course_alignment_score_band_heatmap", run_name, row["course"].id, row["course"].name, row["school"],
                    "", "", "", cell["band"], cell["count"], row["avg_score"], "", "", "", "",
                    "Cell value is count of job adverts in this score band",
                ])
        for point in visual_data["scatter_points"]:
            writer.writerow([
                "matched_vs_missing_scatter", run_name, "", point["course"], point["school"],
                "", "", "", "matched_missing_totals", "", point["score"], point["x"], point["y"], "", "",
                "x=matched skill evidence, y=missing skill evidence",
            ])
        for row in visual_data["skill_suggestion_matrix"]:
            writer.writerow([
                "school_skill_matrix", run_name, "", "", request.GET.get("school", "").strip() or "Current selection",
                "", "", "", row["skill"], row["demand_count"], "", row["matched_count"], row["missing_count"], "", "",
                f"gap_percent={row['gap_percent']}; coverage_percent={row['coverage']}; jobs={row['job_count']}; courses_covered={row['course_count']}",
            ])
        for item in visual_data["school_recommendations"]:
            writer.writerow([
                "curriculum_recommendation", run_name, "", "", item["school"],
                "", "", "", "recommendation", item["message"], item["score"], "", "", "", "",
                "Generated from selected threshold and evidence counts",
            ])
        return response


def docx_add_field(paragraph, instruction, placeholder="Right-click and update field in Word."):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def docx_shade(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def docx_set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold


def add_key_value_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    docx_set_cell_text(hdr[0], "Metric", True)
    docx_set_cell_text(hdr[1], "Value", True)
    docx_shade(hdr[0], "F2EEE8")
    docx_shade(hdr[1], "F2EEE8")
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    return table


def add_simple_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        docx_set_cell_text(table.rows[0].cells[index], header, True)
        docx_shade(table.rows[0].cells[index], "F2EEE8")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def score_fill(score):
    if score >= 80:
        return "F58220"
    if score >= 60:
        return "FFBD80"
    if score >= 40:
        return "FFE1C2"
    if score > 0:
        return "F1E7DD"
    return "F1F3F5"


def add_cross_tab_table(document, results, limit_courses=12, limit_jobs=10):
    courses = []
    jobs = []
    for result in results:
        if result.course not in courses:
            courses.append(result.course)
        if result.job not in jobs:
            jobs.append(result.job)
    courses = courses[:limit_courses]
    jobs = jobs[:limit_jobs]
    if not courses or not jobs:
        document.add_paragraph("No cross-tab data available for this school.")
        return

    score_map = {(result.course_id, result.job_id): result.similarity_percent for result in results}
    headers = ["Course"] + [job.title[:28] for job in jobs]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        docx_set_cell_text(table.rows[0].cells[index], header, True)
        docx_shade(table.rows[0].cells[index], "F2EEE8")
    for course in courses:
        cells = table.add_row().cells
        cells[0].text = course.code or course.name[:32]
        for index, job in enumerate(jobs, start=1):
            score = score_map.get((course.id, job.id), 0)
            cells[index].text = f"{score:.0f}" if score else ""
            docx_shade(cells[index], score_fill(score))


def add_skill_matrix_table(document, matrix_rows):
    if not matrix_rows:
        document.add_paragraph("No school skill matrix data available.")
        return
    rows = [
        [
            row["skill"],
            row["demand_count"],
            row["job_count"],
            row["course_count"],
            row["missing_count"],
            f'{row["gap_percent"]}%',
            f'{row["coverage"]}%',
        ]
        for row in matrix_rows
    ]
    add_simple_table(
        document,
        ["Skill", "Demand evidence", "Jobs", "Courses covered", "Gap evidence", "Gap %", "Coverage %"],
        rows,
    )


def plotly_methodology_image(kind):
    try:
        import plotly.graph_objects as go
    except Exception:
        return None

    palette = {
        "black": "#0b0b0b",
        "gold": "#f58220",
        "blue": "#2f80ed",
        "green": "#236b35",
        "muted": "#706960",
    }
    if kind == "pipeline":
        labels = ["Ingest", "Clean", "Extract", "Compare", "Analyse", "Export"]
        details = [
            "courses, uploads, jobs",
            "normalised text, fingerprints",
            "NER, aliases, sectors",
            "vectors, cosine, coverage",
            "matched and missing skills",
            "heatmaps, clusters, CSV",
        ]
        x_values = list(range(len(labels)))
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=x_values,
            y=[0] * len(labels),
            mode="lines+markers+text",
            line={"color": palette["gold"], "width": 5},
            marker={"size": 34, "color": [palette["black"], palette["muted"], palette["gold"], palette["blue"], palette["green"], palette["black"]]},
            text=[str(index + 1) for index in x_values],
            textfont={"color": "white", "size": 13},
            textposition="middle center",
            hoverinfo="skip",
            showlegend=False,
        ))
        for index, (label, detail) in enumerate(zip(labels, details)):
            fig.add_annotation(x=index, y=.42, text=f"<b>{label}</b><br>{detail}", showarrow=False, align="center", font={"size": 13})
            if index < len(labels) - 1:
                fig.add_annotation(x=index + .5, y=.02, ax=index + .28, ay=.02, xref="x", yref="y", axref="x", ayref="y", showarrow=True, arrowhead=3, arrowsize=1.1, arrowwidth=2, arrowcolor=palette["gold"], text="")
        fig.update_layout(
            title={"text": "Methodology Pipeline", "x": .02, "xanchor": "left"},
            width=1050,
            height=360,
            margin={"l": 20, "r": 20, "t": 60, "b": 30},
            plot_bgcolor="white",
            paper_bgcolor="white",
            xaxis={"visible": False, "range": [-.5, len(labels) - .5]},
            yaxis={"visible": False, "range": [-.45, .78]},
        )
    elif kind == "funnel":
        fig = go.Figure(go.Funnel(
            y=["Raw evidence", "Structured records", "Skill evidence", "Comparable signals", "Actionable outputs"],
            x=[100, 82, 62, 42, 25],
            text=[
                "courses, jobs, files, scraped pages",
                "modules, adverts, dates, sectors",
                "NER entities, skill types, confidence",
                "embeddings, cosine, skill coverage",
                "gaps, forecasts, clusters, reports",
            ],
            textposition="inside",
            marker={"color": [palette["black"], "#3c3327", "#7b4b16", palette["gold"], palette["green"]]},
            connector={"line": {"color": palette["muted"], "width": 1}},
            hoverinfo="skip",
        ))
        fig.update_layout(
            title={"text": "Evidence Funnel", "x": .02, "xanchor": "left"},
            width=850,
            height=430,
            margin={"l": 30, "r": 30, "t": 60, "b": 20},
            paper_bgcolor="white",
            font={"size": 13},
        )
    else:
        fig = go.Figure()
        fig.add_trace(go.Bar(
            x=["Semantic similarity", "Skill coverage", "Review status", "Forecast trend", "Vector space"],
            y=[75, 25, 100, 60, 80],
            marker={"color": [palette["black"], palette["gold"], palette["green"], palette["blue"], palette["muted"]]},
            text=["75% final-score weight", "25% final-score weight", "auditable evidence", "yearly skill counts", "skill associations"],
            textposition="outside",
            hoverinfo="skip",
            showlegend=False,
        ))
        for name, color in [
            ("Semantic", palette["black"]),
            ("Skill evidence", palette["gold"]),
            ("Reviewable", palette["green"]),
            ("Trend", palette["blue"]),
            ("Network", palette["muted"]),
        ]:
            fig.add_trace(go.Scatter(x=[None], y=[None], mode="markers", marker={"size": 12, "color": color}, name=name))
        fig.update_layout(
            title={"text": "Report Graph Legend and Outputs", "x": .02, "xanchor": "left"},
            width=900,
            height=420,
            margin={"l": 55, "r": 35, "t": 60, "b": 90},
            plot_bgcolor="white",
            paper_bgcolor="white",
            yaxis={"title": "Relative role in methodology", "range": [0, 115], "gridcolor": "#f2eee8"},
            xaxis={"tickangle": -18},
            legend={"orientation": "h", "y": -0.28},
        )

    try:
        return BytesIO(fig.to_image(format="png", scale=2))
    except Exception:
        return None


def add_report_visual_fallback(document, title, rows):
    document.add_paragraph(title)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, header in enumerate(["Stage", "What happens", "Output"]):
        docx_set_cell_text(table.rows[0].cells[index], header, True)
        docx_shade(table.rows[0].cells[index], "F2EEE8")
    for stage, action, output in rows:
        cells = table.add_row().cells
        cells[0].text = stage
        cells[1].text = action
        cells[2].text = output


def add_methodology_visuals_to_docx(document):
    from docx.shared import Inches

    document.add_heading("Methodology Visual Pipeline", level=2)
    document.add_paragraph(
        "The figures below summarise the end-to-end flow from raw curriculum and labour-market evidence into reviewable skill evidence, match scores, visual analytics, and exports."
    )
    visuals = [
        (
            "pipeline",
            "Figure 1. End-to-end methodology pipeline.",
            [
                ("1. Ingest", "Collect courses, modules, uploads, scraped pages, CSV/manual jobs, and Adzuna adverts.", "Raw evidence store"),
                ("2. Clean", "Normalise text, parse job sections, handle stale tasks, and skip duplicate adverts.", "Structured records"),
                ("3. Extract", "Apply spaCy NER, phrase matching, aliases, regex fallback, and phrase mining.", "Skill entities"),
                ("4. Compare", "Create normalised embeddings and compute cosine similarity plus skill coverage.", "Course-to-job scores"),
                ("5. Analyse", "Calculate matched skills, missing skills, score bands, sectors, and school matrices.", "Gap evidence"),
                ("6. Export", "Render dashboards, heatmaps, forecasts, clusters, vector space, CSV, and report outputs.", "Decision support"),
            ],
        ),
        (
            "funnel",
            "Figure 2. Evidence funnel.",
            [
                ("Raw evidence", "Courses, files, jobs, and scraped pages enter the pipeline.", "Input corpus"),
                ("Structured records", "Text is linked to modules, adverts, dates, titles, and sectors.", "Comparable records"),
                ("Skill evidence", "NER entities store skill, type, source, confidence, and review status.", "Auditable skills"),
                ("Comparable signals", "Embeddings and skill coverage convert text into match signals.", "Scores"),
                ("Actionable outputs", "Gaps, forecasts, clusters, exports, and reports support decisions.", "Insights"),
            ],
        ),
        (
            "legend",
            "Figure 3. High-level graph legend.",
            [
                ("Semantic similarity", "Normalised vector cosine similarity contributes the main score signal.", "Black"),
                ("Skill evidence", "Extracted skills and coverage contribute explicit evidence.", "Gold"),
                ("Reviewable records", "Inline review updates stored job/module skill entities.", "Green"),
                ("Trend", "Forecasts use extraction/posting year counts.", "Blue"),
                ("Network", "Vector space and clusters show skill associations.", "Grey"),
            ],
        ),
    ]
    for kind, caption, fallback_rows in visuals:
        image = plotly_methodology_image(kind)
        if image:
            document.add_picture(image, width=Inches(6.4))
            document.add_paragraph(caption)
        else:
            add_report_visual_fallback(document, caption, fallback_rows)


def latest_visual_run():
    run = AnalysisRun.objects.first()
    if run and not GapResult.objects.filter(run=run).exists():
        run = AnalysisRun.objects.filter(results__isnull=False).distinct().order_by("-created_at").first()
    return run


class DashboardVisualCsvExportView(View):
    def get(self, request):
        visual_run = latest_visual_run()
        results = list(GapResult.objects.filter(run=visual_run).select_related("course", "job")) if visual_run else []
        response = csv_download_response("dashboard-visual-source.csv")
        writer = csv.writer(response)
        writer.writerow([
            "visual", "run", "source", "course_id", "course", "job_id", "job",
            "dimension", "value", "score_percent", "matched_skills", "missing_skills", "notes",
        ])
        run_name = visual_run.name if visual_run else ""

        score_labels = ["0-20", "20-40", "40-60", "60-80", "80-100"]
        buckets = [0, 0, 0, 0, 0]
        for result in results:
            buckets[min(4, int(max(0, result.similarity_score) * 5))] += 1
        for label, count in zip(score_labels, buckets):
            writer.writerow(["score_bucket_summary", run_name, "analysis", "", "", "", "", label, count, "", "", "", "Count of course-job scores in this range"])

        if visual_run:
            for source in ["jobs", "courses"]:
                for row in SkillMatrix.objects.filter(run=visual_run, source=source).values("skill", "frequency")[:50]:
                    writer.writerow(["skill_demand_vs_curriculum", run_name, source, "", "", "", "", row["skill"], row["frequency"], "", "", "", "Skill frequency shown in dashboard bar/line chart"])

        for result in sorted(results, key=lambda item: item.similarity_score, reverse=True)[:100]:
            writer.writerow([
                "course_to_job_network_edge", run_name, "similarity", result.course_id, result.course.name,
                result.job_id, result.job.title, "cosine_similarity", result.similarity_score,
                result.similarity_percent, "; ".join(result.matched_skills or []), "; ".join(result.missing_skills or []),
                "Edge/node source for dashboard similarity network and best-match tiles",
            ])
        return response


class TechnicalReportExportView(View):
    def get(self, request):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return HttpResponse("python-docx is not installed.", status=500)

        run = latest_visual_run()
        schools = list(
            Course.objects
            .exclude(university_name="")
            .order_by("university_name")
            .values_list("university_name", flat=True)
            .distinct()
        )
        if Course.objects.filter(university_name="").exists():
            schools.append("Unassigned school")

        document = Document()
        section = document.sections[0]
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        docx_add_field(footer, "PAGE", "1")

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("CurriculumMatch Technical Report")
        title_run.bold = True
        title_run.font.size = None
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("Curriculum-to-job-market alignment report").italic = True
        document.add_paragraph(f"Generated: {timezone.localtime(timezone.now()).strftime('%d %B %Y %H:%M')}")
        document.add_paragraph(f"Analysis run: {run.name if run else 'No completed analysis run available'}")
        document.add_paragraph("Scope: all schools currently stored in the database.")
        document.add_page_break()

        document.add_heading("Table of Contents", level=1)
        docx_add_field(document.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')
        document.add_paragraph("Note: in Microsoft Word, right-click the table above and choose Update Field to refresh page numbers.")
        document.add_page_break()

        document.add_heading("1. Introduction", level=1)
        document.add_paragraph(
            "This technical report summarises curriculum-to-job-market alignment using the courses, modules, "
            "job adverts, extracted skills, semantic similarity scores, and school metadata currently stored in CurriculumMatch. "
            "It is designed as an operational report, not a literature review."
        )

        document.add_heading("2. Methodology", level=1)
        document.add_paragraph(
            "The system parses curriculum and job evidence, extracts explicit skill terms, creates semantic vectors, "
            "and compares each course against job adverts. The final score blends semantic similarity with explicit skill coverage."
        )
        document.add_paragraph("Cosine similarity: cos(A, B) = (A . B) / (||A|| * ||B||)")
        document.add_paragraph("Course semantic score: mean(top-k module-to-job cosine scores)")
        document.add_paragraph("Skill coverage: matched job skills / unique job skills")
        document.add_paragraph("Final score: ((0.75 * semantic) + (0.25 * skill coverage)) / (0.75 + 0.25)")
        document.add_paragraph(
            "School Skill Matrix: demand evidence is the count of course-to-job comparisons where a refined skill appears as "
            "matched or missing. Gap percentage is gap evidence divided by demand evidence. Coverage percentage is covered "
            "evidence divided by demand evidence."
        )
        add_methodology_visuals_to_docx(document)

        document.add_heading("3. Dashboard Snapshot", level=1)
        all_results = list(GapResult.objects.filter(run=run).select_related("course", "job")) if run else []
        avg_score = round((sum(result.similarity_score for result in all_results) / len(all_results)) * 100, 1) if all_results else 0
        add_key_value_table(document, [
            ("Schools", len(schools)),
            ("Courses", Course.objects.count()),
            ("Job adverts", JobAdvert.objects.count()),
            ("Course-job comparisons", len(all_results)),
            ("Average final score", f"{avg_score}%"),
        ])

        document.add_heading("4. School Results", level=1)
        for school in schools:
            document.add_heading(school, level=2)
            if school == "Unassigned school":
                school_results = [result for result in all_results if not result.course.university_name]
            else:
                school_results = [result for result in all_results if result.course.university_name == school]
            school_courses = Course.objects.filter(university_name="" if school == "Unassigned school" else school)
            visual_data = build_results_visual_data(school_results, 55)
            summary = visual_data["school_summaries"][0] if visual_data["school_summaries"] else None
            add_key_value_table(document, [
                ("Courses in database", school_courses.count()),
                ("Analysed comparisons", len(school_results)),
                ("Average score", f'{summary["avg_score"]}%' if summary else "No analysis data"),
                ("Matched evidence", summary["matched_total"] if summary else 0),
                ("Missing evidence", summary["missing_total"] if summary else 0),
                ("Mismatch risk", "Yes" if summary and summary["mismatch"] else "No"),
            ])

            document.add_heading("Curriculum Recommendations", level=3)
            for item in visual_data["school_recommendations"]:
                document.add_paragraph(f'{item["school"]} ({item["score"]}%): {item["message"]}', style=None)

            document.add_heading("School Skill Matrix", level=3)
            add_skill_matrix_table(document, visual_data["skill_suggestion_matrix"])

            document.add_heading("Course-to-Job Cross-tab Snapshot", level=3)
            add_cross_tab_table(document, school_results)

            document.add_heading("Top Course-to-Job Matches", level=3)
            top_rows = [
                [
                    result.course.name[:48],
                    result.job.title[:48],
                    result.job.company or "",
                    f"{result.similarity_percent}%",
                    ", ".join((result.matched_skills or [])[:5]),
                    ", ".join((result.missing_skills or [])[:5]),
                ]
                for result in sorted(school_results, key=lambda item: item.similarity_score, reverse=True)[:8]
            ]
            if top_rows:
                add_simple_table(document, ["Course", "Job advert", "Company", "Score", "Matched", "Missing"], top_rows)
            else:
                document.add_paragraph("No analysed matches available for this school.")

        document.add_heading("5. Appendix: Interpretation Notes", level=1)
        document.add_paragraph(
            "Scores and evidence counts are decision-support signals. Low alignment can indicate a genuine curriculum gap, "
            "a job advert outside programme scope, thin module text, or terminology differences. Recommendations should be "
            "reviewed by curriculum owners before module changes are made."
        )

        output = BytesIO()
        document.save(output)
        output.seek(0)
        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="curriculummatch-technical-report.docx"'
        return response


class TaskListView(ListView):
    model = TaskRecord
    template_name = "analysis/tasks.html"
    context_object_name = "tasks"
    paginate_by = 40

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx["has_running"] = TaskRecord.objects.filter(status__in=["PENDING", "STARTED"]).exists()
        return ctx


def task_status_api(request, pk):
    r = get_object_or_404(TaskRecord, pk=pk)
    r = mark_stale_task_if_needed(r)
    return JsonResponse({
        "id": r.id,
        "run_name": r.run_name,
        "status": r.status,
        "progress": r.progress,
        "notes": r.notes,
        "debug_hint": task_debug_hint(r.notes) if r.status == "FAILURE" else "",
        "updated_at": r.updated_at.isoformat() if r.updated_at else None,
        "finished_at": r.finished_at.isoformat() if r.finished_at else None,
    })


def dashboard_metrics(request):
    last_run = AnalysisRun.objects.first()
    visual_run = last_run
    if visual_run and not GapResult.objects.filter(run=visual_run).exists():
        visual_run = AnalysisRun.objects.filter(results__isnull=False).distinct().order_by("-created_at").first()
    results = GapResult.objects.filter(run=visual_run).select_related("course", "job") if visual_run else GapResult.objects.none()
    score_values = list(results.values_list("similarity_score", flat=True))
    buckets = [0, 0, 0, 0, 0]
    for score in score_values:
        idx = min(4, int(max(0, score) * 5))
        buckets[idx] += 1

    job_skills = []
    course_skills = []
    if visual_run:
        job_skill_rows = list(SkillMatrix.objects.filter(run=visual_run, source="jobs").values("skill", "frequency")[:30])
        course_skill_rows = list(SkillMatrix.objects.filter(run=visual_run, source="courses").values("skill", "frequency")[:30])
        job_skills = actual_skill_rows(job_skill_rows)
        course_skills = actual_skill_rows(course_skill_rows)

    recent_tasks = []
    for task in TaskRecord.objects.order_by("-created_at").values("id", "run_name", "status", "progress", "notes")[:8]:
        task["debug_hint"] = task_debug_hint(task.get("notes")) if task["status"] == "FAILURE" else ""
        recent_tasks.append(task)
    return JsonResponse({
        "counts": {
            "courses": Course.objects.count(),
            "modules": sum(c.modules.count() for c in Course.objects.prefetch_related("modules")),
            "jobs": JobAdvert.objects.count(),
            "runs": AnalysisRun.objects.count(),
        },
        "last_run": {
            "id": last_run.id,
            "status": last_run.status,
            "name": last_run.name,
        } if last_run else None,
        "visual_run": {
            "id": visual_run.id,
            "status": visual_run.status,
            "name": visual_run.name,
        } if visual_run else None,
        "has_visual_data": bool(score_values or job_skills or course_skills),
        "average_score": round((sum(score_values) / len(score_values)) * 100, 1) if score_values else 0,
        "score_buckets": buckets,
        "score_labels": ["0-20", "20-40", "40-60", "60-80", "80-100"],
        "job_skills": job_skills,
        "course_skills": course_skills,
        "recent_tasks": recent_tasks,
    })


def similarity_network(request):
    try:
        import networkx as nx
    except ImportError:
        nx = None

    last_run = AnalysisRun.objects.first()
    visual_run = last_run
    if visual_run and not GapResult.objects.filter(run=visual_run).exists():
        visual_run = AnalysisRun.objects.filter(results__isnull=False).distinct().order_by("-created_at").first()
    school_filter = request.GET.get("school", "").strip()
    job_filter = bounded_int(request.GET.get("job"), 0, 0, 100000000) or None
    include_skills = request.GET.get("cluster") == "1"
    qs = GapResult.objects.none()
    if visual_run:
        qs = GapResult.objects.filter(run=visual_run).select_related("course", "job")
        if school_filter:
            qs = qs.filter(course__university_name=school_filter)
        if job_filter:
            qs = qs.filter(job_id=job_filter)
        qs = qs.order_by("-similarity_score")[:40]

    def job_node_title(result, skills):
        title = result.job.title
        if result.job.company:
            title = f"{title} @ {result.job.company}"
        if skills:
            title += "\nSkills: " + ", ".join(entity["skill"] for entity in skills[:10])
        return title

    def add_cluster_skill_nodes(graph, result, job_id, skills):
        for entity in skills:
            skill_id = entity["id"]
            graph.add_node(
                skill_id,
                label=entity["skill"][:26],
                group="skill",
                title=f"Skill: {entity['skill']}\nEntity ID: {skill_id}\nSource: {entity['source']}",
                skill=entity["skill"],
                entity_id=skill_id,
                source=entity["source"],
                confidence=entity["confidence"],
                mention_count=entity["mention_count"],
            )
            graph.add_edge(
                job_id,
                skill_id,
                value=max(8, entity["mention_count"] * 8),
                group="skill-link",
                title=f"{result.job.title} requires {entity['skill']}",
            )

    if nx:
        graph = nx.Graph()
        for r in qs:
            course_id = f"course-{r.course_id}"
            job_id = f"job-{r.job_id}"
            skills = job_skill_entities(r.job, (r.matched_skills or []) + (r.missing_skills or []))
            graph.add_node(course_id, label=r.course.code or r.course.name, group="course", title=r.course.name)
            graph.add_node(
                job_id,
                label=r.job.title[:28],
                group="job",
                title=job_node_title(r, skills),
                skills=[entity["skill"] for entity in skills],
                skill_entities=skills,
            )
            graph.add_edge(course_id, job_id, value=max(1, r.similarity_percent), title=f"Cosine similarity: {r.similarity_score:.4f}")
            if include_skills:
                add_cluster_skill_nodes(graph, r, job_id, skills)
        positions = nx.spring_layout(graph, seed=42, k=0.95 if include_skills else 0.7) if graph.number_of_nodes() else {}
        nodes = [
            {
                "id": n,
                "x": round(positions.get(n, (0, 0))[0] * (700 if include_skills else 520), 2),
                "y": round(positions.get(n, (0, 0))[1] * (460 if include_skills else 360), 2),
                **attrs,
            }
            for n, attrs in graph.nodes(data=True)
        ]
        edges = [{"from": u, "to": v, **attrs} for u, v, attrs in graph.edges(data=True)]
    else:
        nodes, edges, seen = [], [], set()
        for index, r in enumerate(qs):
            course_id = f"course-{r.course_id}"
            job_id = f"job-{r.job_id}"
            skills = job_skill_entities(r.job, (r.matched_skills or []) + (r.missing_skills or []))
            if course_id not in seen:
                nodes.append({"id": course_id, "label": r.course.code or r.course.name, "group": "course", "title": r.course.name, "x": -260, "y": index * 52})
                seen.add(course_id)
            if job_id not in seen:
                nodes.append({
                    "id": job_id,
                    "label": r.job.title[:28],
                    "group": "job",
                    "title": job_node_title(r, skills),
                    "skills": [entity["skill"] for entity in skills],
                    "skill_entities": skills,
                    "x": 260,
                    "y": index * 52,
                })
                seen.add(job_id)
            edges.append({"from": course_id, "to": job_id, "value": max(1, r.similarity_percent), "title": f"Cosine similarity: {r.similarity_score:.4f}"})
            if include_skills:
                for skill_index, entity in enumerate(skills):
                    if entity["id"] not in seen:
                        nodes.append({
                            "id": entity["id"],
                            "label": entity["skill"][:26],
                            "group": "skill",
                            "title": f"Skill: {entity['skill']}\nEntity ID: {entity['id']}\nSource: {entity['source']}",
                            "skill": entity["skill"],
                            "entity_id": entity["id"],
                            "source": entity["source"],
                            "confidence": entity["confidence"],
                            "mention_count": entity["mention_count"],
                            "x": 520,
                            "y": (index * 52) + (skill_index - len(skills) / 2) * 28,
                        })
                        seen.add(entity["id"])
                    edges.append({
                        "from": job_id,
                        "to": entity["id"],
                        "value": max(8, entity["mention_count"] * 8),
                        "group": "skill-link",
                        "title": f"{r.job.title} requires {entity['skill']}",
                    })
    return JsonResponse({
        "run_id": visual_run.id if visual_run else None,
        "cluster": include_skills,
        "has_visual_data": bool(nodes and edges),
        "nodes": nodes,
        "edges": edges,
    })


def build_skill_vector_space_payload(request):
    rows = filtered_skill_entity_rows(request)[:900]
    job_lookup = JobAdvert.objects.in_bulk({
        row["source_id"] for row in rows if row["source_type"] == "job"
    })
    module_lookup = Module.objects.in_bulk({
        row["source_id"] for row in rows if row["source_type"] == "module"
    })
    nodes = {}
    edges = []
    skill_context = {}

    def add_node(node_id, **attrs):
        nodes.setdefault(node_id, {"id": node_id}).update(attrs)

    for row in rows:
        root_id = f"{row['source_type']}-{row['source_id']}"
        root_group = "job-root" if row["source_type"] == "job" else "course-root"
        root_label = row["job_title"] if row["source_type"] == "job" and row["job_title"] else row["source_label"]
        source_obj = job_lookup.get(row["source_id"]) if row["source_type"] == "job" else module_lookup.get(row["source_id"])
        description = getattr(source_obj, "description", "") or getattr(source_obj, "content", "") or ""
        add_node(
            root_id,
            label=root_label[:54],
            full_label=root_label,
            group=root_group,
            source_type=row["source_type"],
            sector=row["sector"],
            parent_label=row["parent_label"],
            description=description[:420],
        )

        skill_id = row["id"]
        add_node(
            skill_id,
            label=row["skill"][:44],
            full_label=row["skill"],
            group="skill",
            skill=row["skill"],
            skill_type=row["skill_type"],
            sector=row["sector"],
            source=row["source"],
            confidence=row["confidence"],
        )
        edges.append({
            "source": root_id,
            "target": skill_id,
            "value": max(1, int(row.get("mention_count") or 1)),
            "group": "evidence",
            "label": "Skill evidence",
            "title": f"{root_label} -> {row['skill']}",
        })
        item = skill_context.setdefault(skill_id, {
            "skill": row["skill"],
            "skill_type": row["skill_type"],
            "sectors": Counter(),
            "sources": set(),
            "count": 0,
        })
        item["sectors"][row["sector"]] += 1
        item["sources"].add(root_id)
        item["count"] += 1

    skill_items = list(skill_context.items())[:140]
    for index, (left_id, left) in enumerate(skill_items):
        for right_id, right in skill_items[index + 1:]:
            score = token_cosine_similarity(left["skill"], right["skill"])
            if left["skill_type"] and left["skill_type"] == right["skill_type"]:
                score += 0.08
            if (set(left["sectors"]) & set(right["sectors"])):
                score += 0.06
            if score < 0.26:
                continue
            percent = min(100, round(score * 100))
            edges.append({
                "source": left_id,
                "target": right_id,
                "value": max(1, percent / 18),
                "group": "similarity",
                "similarity": percent,
                "label": f"{percent}% cosine",
                "title": f"Cosine similarity: {percent}%\n{left['skill']} <-> {right['skill']}",
            })

    return {
        "has_visual_data": bool(nodes and edges),
        "nodes": list(nodes.values()),
        "edges": edges,
        "counts": {
            "nodes": len(nodes),
            "edges": len(edges),
            "skills": len(skill_context),
            "roots": len([node for node in nodes.values() if node.get("group") != "skill"]),
        },
    }


class SkillVectorSpaceCsvExportView(View):
    def get(self, request):
        payload = build_skill_vector_space_payload(request)
        response = csv_download_response("skill-vector-space-source.csv")
        writer = csv.writer(response)
        writer.writerow([
            "record_type", "id", "source", "target", "label", "group", "value",
            "similarity_percent", "skill", "skill_type", "sector", "source_type", "description",
        ])
        for node in payload["nodes"]:
            writer.writerow([
                "node", node.get("id", ""), "", "", node.get("full_label") or node.get("label", ""),
                node.get("group", ""), "", "", node.get("skill", ""), node.get("skill_type", ""),
                node.get("sector", ""), node.get("source_type", ""), node.get("description", ""),
            ])
        for edge in payload["edges"]:
            writer.writerow([
                "edge", "", edge.get("source", ""), edge.get("target", ""), edge.get("label", ""),
                edge.get("group", ""), edge.get("value", ""), edge.get("similarity", ""),
                "", "", "", "", edge.get("title", ""),
            ])
        return response


def skill_vector_space(request):
    return JsonResponse(build_skill_vector_space_payload(request))


def results_json(request):
    run_id = request.GET.get("run")
    if not run_id:
        return JsonResponse({"error": "run param required"}, status=400)
    qs = GapResult.objects.filter(run_id=run_id).select_related("course", "job")
    return JsonResponse({"results": [
        {
            "course": r.course.name,
            "job": r.job.title,
            "company": r.job.company,
            "recruiter": r.job.recruiter,
            "job_reference": r.job.job_reference,
            "location": r.job.location,
            "category": r.job.category,
            "contract_type": r.job.contract_type,
            "contract_time": r.job.contract_time,
            "date_posted": r.job.date_posted.isoformat() if r.job.date_posted else None,
            "summary": r.job.summary,
            "position_info": r.job.position_info,
            "score": round(r.similarity_score * 100, 1),
            "matched": r.matched_skills,
            "missing": r.missing_skills,
        } for r in qs
    ]})
