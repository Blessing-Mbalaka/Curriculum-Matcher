from collections import Counter
from io import BytesIO
from pathlib import Path
import re

from django.conf import settings
from django.utils import timezone

from analysis.models import GapResult, SkillAlias, SkillMatrix
from courses.models import Course, Module
from jobs.models import JobAdvert


REPORT_FILENAME = "curriculummatch-research-paper.docx"
PAPER_ARTIFACT_ROOT = "paper_artifacts"


def safe_filename(value, fallback="run"):
    name = re.sub(r"[^a-zA-Z0-9._-]+", "-", str(value or "").strip().lower()).strip("-")
    return name[:80] or fallback


def paper_artifact_dir(run):
    run_label = f"{run.id}-{run.name}" if run else "no-run"
    root = Path(settings.BASE_DIR) / PAPER_ARTIFACT_ROOT / safe_filename(run_label) / "images"
    root.mkdir(parents=True, exist_ok=True)
    return root


def write_visual_manifest(asset_dir, rows):
    manifest = asset_dir.parent / "visual_manifest.md"
    lines = [
        "# CurriculumMatch Paper Visual Manifest",
        "",
        f"Generated: {timezone.localtime(timezone.now()).isoformat()}",
        f"Image folder: `{asset_dir}`",
        "",
        "| Figure | File | Caption |",
        "| --- | --- | --- |",
    ]
    for row in rows:
        lines.append(f"| {row['title']} | `{row['file']}` | {row['caption']} |")
    manifest.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return manifest


def docx_add_field(paragraph, instruction, placeholder="Right-click and update field in Word."):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def docx_add_equation(document, equation_text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = equation_text
    run.append(text)
    math.append(run)
    math_para.append(math)
    paragraph._p.append(math_para)
    return paragraph


def docx_shade(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def docx_set_cell_text(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold


def add_simple_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        docx_set_cell_text(table.rows[0].cells[index], header, True)
        docx_shade(table.rows[0].cells[index], "F2EEE8")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def add_key_value_table(document, rows):
    return add_simple_table(document, ["Metric", "Value"], rows)


def plotly_to_image(fig, width=1050, height=560):
    if fig is None:
        return None
    try:
        return BytesIO(fig.to_image(format="png", width=width, height=height, scale=2))
    except Exception:
        return None


def plotly_go():
    try:
        import plotly.graph_objects as go
        return go
    except Exception:
        return None


def fallback_png(title, caption, fallback_headers=None, fallback_rows=None, width=1400, height=760):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None

    def number(value, default=0):
        try:
            return float(str(value).replace("%", "").replace(",", "").strip())
        except (TypeError, ValueError):
            return default

    def wrap_text(text, max_chars=34):
        words = str(text or "").split()
        lines = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if len(candidate) > max_chars and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        return lines or [""]

    def short_text(text, limit=30):
        text = str(text or "")
        return text if len(text) <= limit else f"{text[:limit - 1]}..."

    def palette(value, maximum, colours=("#fff7ed", "#fdba74", "#f58220", "#111111")):
        if maximum <= 0:
            return colours[0]
        ratio = max(0, min(1, value / maximum))
        return colours[min(len(colours) - 1, int(ratio * (len(colours) - 1)))]

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.load_default(size=32)
    body_font = ImageFont.load_default(size=20)
    small_font = ImageFont.load_default(size=16)
    tiny_font = ImageFont.load_default(size=13)
    x = 44
    y = 38
    draw.text((x, y), title, fill="#101010", font=title_font)
    y += 58
    draw.text((x, y), "Rendered by CurriculumMatch from the same source data used by the web visual.", fill="#9a4a08", font=body_font)
    y += 38

    for line in wrap_text(caption, 112)[:2]:
        draw.text((x, y), line, fill="#333333", font=body_font)
        y += 28
    chart_top = y + 24
    chart_left = 86
    chart_right = width - 64
    chart_bottom = height - 78
    rows = [tuple(row) for row in (fallback_rows or []) if row]
    title_lower = title.lower()
    headers = list(fallback_headers or [])

    def axes(left=chart_left, top=chart_top, right=chart_right, bottom=chart_bottom, max_value=None, x_label="", y_label=""):
        draw.line((left, bottom, right, bottom), fill="#706960", width=2)
        draw.line((left, top, left, bottom), fill="#706960", width=2)
        if max_value is not None:
            for tick_index in range(5):
                ratio = tick_index / 4
                value = max_value * ratio
                y_pos = bottom - int((bottom - top) * ratio)
                draw.line((left - 5, y_pos, left, y_pos), fill="#706960", width=1)
                draw.line((left, y_pos, right, y_pos), fill="#f2eee8", width=1)
                draw.text((left - 50, y_pos - 7), str(round(value, 1)).rstrip("0").rstrip("."), fill="#333333", font=tiny_font)
        if x_label:
            draw.text(((left + right) / 2 - 70, height - 30), x_label, fill="#333333", font=small_font)
        if y_label:
            draw.text((18, top - 22), y_label, fill="#333333", font=small_font)

    def legend(items, anchor_x=None, anchor_y=None):
        if not items:
            return
        lx = anchor_x if anchor_x is not None else chart_right - 260
        ly = anchor_y if anchor_y is not None else chart_top
        for index, (label, colour) in enumerate(items):
            item_y = ly + index * 24
            draw.rectangle((lx, item_y, lx + 14, item_y + 14), fill=colour)
            draw.text((lx + 20, item_y - 1), str(label), fill="#333333", font=small_font)

    def draw_bars(grouped=False):
        if not rows:
            draw.text((chart_left, chart_top), "No chart data available.", fill="#706960", font=body_font)
            return
        labels = [short_text(row[0], 24) for row in rows[:18]]
        series_one = [number(row[1]) if len(row) > 1 else 0 for row in rows[:18]]
        series_two = [number(row[2]) if grouped and len(row) > 2 else 0 for row in rows[:18]]
        series_one = [max(0, value) for value in series_one]
        series_two = [max(0, value) for value in series_two]
        max_value = max(series_one + series_two + [1])
        y_label = headers[1] if len(headers) > 1 else "Value"
        axes(max_value=max_value, x_label=headers[0] if headers else "", y_label="Records" if "record" in str(y_label).lower() else str(y_label))
        bar_gap = 10
        usable_width = chart_right - chart_left - 24
        slot = max(18, usable_width / max(1, len(labels)))
        bar_width = max(8, min(42, (slot - bar_gap) / (2 if grouped else 1)))
        for index, label in enumerate(labels):
            base_x = chart_left + 18 + index * slot
            values = [series_one[index]] + ([series_two[index]] if grouped else [])
            colours = ["#0b0b0b", "#f58220"] if grouped else ["#f58220"]
            for series_index, value in enumerate(values):
                left = base_x + series_index * (bar_width + 3)
                bar_height = int((chart_bottom - chart_top - 34) * value / max_value)
                draw.rectangle((left, chart_bottom - bar_height, left + bar_width, chart_bottom), fill=colours[series_index])
                draw.text((left, chart_bottom - bar_height - 18), str(int(value)), fill="#333333", font=tiny_font)
            for line_index, line in enumerate(wrap_text(label, 12)[:2]):
                draw.text((base_x - 4, chart_bottom + 8 + line_index * 14), line, fill="#333333", font=tiny_font)
        if grouped:
            legend([
                ((headers or ["", "Series 1"])[1], "#0b0b0b"),
                ((headers or ["", "", "Series 2"])[2], "#f58220"),
            ])

    def draw_heatmap():
        if not rows:
            draw.text((chart_left, chart_top), "No heatmap data available.", fill="#706960", font=body_font)
            return
        row_labels = []
        col_labels = []
        values = {}
        for row in rows[:260]:
            left = str(row[0])
            col = str(row[1]) if len(row) > 1 else ""
            value = number(row[2] if len(row) > 2 else 0)
            if left not in row_labels:
                row_labels.append(left)
            if col not in col_labels:
                col_labels.append(col)
            values[(left, col)] = value
        row_labels = row_labels[:18]
        col_labels = col_labels[:24]
        max_value = max(values.values() or [1])
        draw.text((chart_left, chart_top - 22), f"Y: {headers[0] if headers else 'Rows'}", fill="#333333", font=small_font)
        draw.text((chart_left + 360, chart_top - 22), f"X: {headers[1] if len(headers) > 1 else 'Columns'}", fill="#333333", font=small_font)
        label_w = 230
        top_label_h = 64
        cell_w = max(30, min(78, (chart_right - chart_left - label_w) / max(1, len(col_labels))))
        cell_h = max(24, min(42, (chart_bottom - chart_top - top_label_h) / max(1, len(row_labels))))
        for c_index, col in enumerate(col_labels):
            cx = chart_left + label_w + c_index * cell_w
            draw.text((cx, chart_top + 18), short_text(col, 13), fill="#333333", font=tiny_font)
        for r_index, row_label in enumerate(row_labels):
            cy = chart_top + top_label_h + r_index * cell_h
            draw.text((chart_left, cy + 7), short_text(row_label, 28), fill="#333333", font=tiny_font)
            for c_index, col in enumerate(col_labels):
                value = values.get((row_label, col), 0)
                cx = chart_left + label_w + c_index * cell_w
                fill = palette(value, max_value)
                draw.rectangle((cx, cy, cx + cell_w - 3, cy + cell_h - 3), fill=fill, outline="#ffffff")
                if cell_w >= 38 and cell_h >= 28:
                    draw.text((cx + 6, cy + 6), str(round(value, 1)).rstrip("0").rstrip("."), fill="#111111", font=tiny_font)
        legend_x = chart_right - 180
        legend_y = chart_bottom - 112
        for index, colour in enumerate(("#fff7ed", "#fdba74", "#f58220", "#111111")):
            draw.rectangle((legend_x + index * 32, legend_y, legend_x + 30 + index * 32, legend_y + 14), fill=colour)
        draw.text((legend_x, legend_y + 20), f"0 to {round(max_value, 1)}", fill="#333333", font=tiny_font)

    def draw_line():
        if not rows:
            draw.text((chart_left, chart_top), "No line chart data available.", fill="#706960", font=body_font)
            return
        labels = [short_text(row[0], 18) for row in rows[:16]]
        series = [number(row[1]) if len(row) > 1 else 0 for row in rows[:16]]
        second = [number(row[2]) if len(row) > 2 else None for row in rows[:16]]
        max_value = max([value for value in series + [v for v in second if v is not None] if value is not None] + [1])
        axes(max_value=max_value, x_label=headers[0] if headers else "", y_label="Records")
        def points(values):
            pts = []
            for index, value in enumerate(values):
                if value is None:
                    continue
                px = chart_left + 24 + index * ((chart_right - chart_left - 52) / max(1, len(labels) - 1))
                py = chart_bottom - int((chart_bottom - chart_top - 30) * value / max_value)
                pts.append((px, py))
            return pts
        for pts, colour in ((points(series), "#0b0b0b"), (points(second), "#f58220")):
            if len(pts) > 1:
                draw.line(pts, fill=colour, width=4)
            for px, py in pts:
                draw.ellipse((px - 5, py - 5, px + 5, py + 5), fill=colour)
        for index, label in enumerate(labels):
            px = chart_left + 24 + index * ((chart_right - chart_left - 52) / max(1, len(labels) - 1))
            draw.text((px - 24, chart_bottom + 10), label, fill="#333333", font=tiny_font)
        if second and any(value is not None for value in second):
            legend([
                ((headers or ["", "Series 1"])[1], "#0b0b0b"),
                ((headers or ["", "", "Series 2"])[2], "#f58220"),
            ])

    def draw_scatter():
        if not rows:
            draw.text((chart_left, chart_top), "No scatter data available.", fill="#706960", font=body_font)
            return
        x_values = [number(row[1]) if len(row) > 1 else 0 for row in rows[:80]]
        y_values = [number(row[2]) if len(row) > 2 else 0 for row in rows[:80]]
        max_x = max(x_values + [1])
        max_y = max(y_values + [1])
        axes(max_value=max_y, x_label=headers[1] if len(headers) > 1 else "X value", y_label=headers[2] if len(headers) > 2 else "Y value")
        for tick_index in range(5):
            ratio = tick_index / 4
            x_pos = chart_left + int((chart_right - chart_left) * ratio)
            draw.line((x_pos, chart_bottom, x_pos, chart_bottom + 5), fill="#706960", width=1)
            draw.text((x_pos - 8, chart_bottom + 10), str(round(max_x * ratio, 1)).rstrip("0").rstrip("."), fill="#333333", font=tiny_font)
        for row, x_value, y_value in zip(rows[:80], x_values, y_values):
            px = chart_left + 18 + int((chart_right - chart_left - 42) * x_value / max_x)
            py = chart_bottom - 18 - int((chart_bottom - chart_top - 42) * y_value / max_y)
            colour = "#a33a2d" if (len(row) > 3 and number(row[3]) < 55) else "#f58220"
            draw.ellipse((px - 8, py - 8, px + 8, py + 8), fill=colour, outline="#111111", width=1)
            draw.text((px + 10, py - 8), short_text(row[0], 18), fill="#333333", font=tiny_font)
        legend([("Below threshold", "#a33a2d"), ("At/above threshold", "#f58220")])

    def draw_diverging_bars():
        if not rows:
            draw.text((chart_left, chart_top), "No divergence data available.", fill="#706960", font=body_font)
            return
        labels = []
        values = []
        for row in rows[:18]:
            labels.append(short_text(row[0] if len(row) == 2 else row[1], 32))
            values.append(number(row[-1]))
        max_abs = max([abs(value) for value in values] + [0.1])
        left = chart_left + 250
        right = chart_right - 40
        zero_x = (left + right) / 2
        draw.line((left, chart_bottom, right, chart_bottom), fill="#706960", width=2)
        draw.line((zero_x, chart_top, zero_x, chart_bottom), fill="#333333", width=2)
        for tick in (-max_abs, -max_abs / 2, 0, max_abs / 2, max_abs):
            x_pos = zero_x + (tick / max_abs) * ((right - left) / 2)
            draw.line((x_pos, chart_bottom, x_pos, chart_bottom + 5), fill="#706960", width=1)
            draw.text((x_pos - 14, chart_bottom + 10), f"{tick:.2f}", fill="#333333", font=tiny_font)
        slot = max(20, (chart_bottom - chart_top - 30) / max(1, len(labels)))
        bar_height = min(24, slot * .62)
        for index, (label_text, value) in enumerate(zip(labels, values)):
            y_mid = chart_top + 18 + index * slot
            draw.text((chart_left, y_mid - 7), label_text, fill="#333333", font=tiny_font)
            x_end = zero_x + (value / max_abs) * ((right - left) / 2)
            colour = "#c83a2f" if value >= 0 else "#3f6fa9"
            draw.rectangle((min(zero_x, x_end), y_mid - bar_height / 2, max(zero_x, x_end), y_mid + bar_height / 2), fill=colour)
            draw.text((x_end + (6 if value >= 0 else -42), y_mid - 7), f"{value:.2f}", fill="#333333", font=tiny_font)
        draw.text((left, height - 30), headers[-1] if headers else "Divergence value", fill="#333333", font=small_font)
        legend([("Above baseline", "#c83a2f"), ("Below baseline", "#3f6fa9")], chart_right - 230, chart_top)

    def draw_network():
        import math
        labels = []
        for row in rows[:48]:
            for value in row[:2]:
                text = short_text(value, 22)
                if text and text not in labels:
                    labels.append(text)
        if not labels:
            labels = [short_text(row[0], 22) for row in rows[:36]]
        if not labels:
            draw.text((chart_left, chart_top), "No network data available.", fill="#706960", font=body_font)
            return
        center_x = (chart_left + chart_right) / 2
        center_y = (chart_top + chart_bottom) / 2 + 20
        radius = min(chart_right - chart_left, chart_bottom - chart_top) * .38
        positions = {}
        for index, label in enumerate(labels[:42]):
            angle = (math.pi * 2 * index / max(1, min(42, len(labels)))) - math.pi / 2
            px = center_x + math.cos(angle) * radius
            py = center_y + math.sin(angle) * radius
            positions[label] = (px, py)
        for row in rows[:80]:
            if len(row) < 2:
                continue
            left = short_text(row[0], 22)
            right = short_text(row[1], 22)
            if left in positions and right in positions:
                draw.line((*positions[left], *positions[right]), fill="#d47722", width=2)
        for index, (label, (px, py)) in enumerate(positions.items()):
            colour = "#111111" if index % 3 == 0 else "#f58220" if index % 3 == 1 else "#2f80ed"
            draw.ellipse((px - 13, py - 13, px + 13, py + 13), fill=colour, outline="#ffffff", width=2)
            draw.text((px + 16, py - 7), label, fill="#333333", font=tiny_font)
        legend([("Course/root", "#111111"), ("Job/skill", "#f58220"), ("Skill/source", "#2f80ed"), ("Relationship", "#d47722")], chart_left, chart_bottom - 92)

    def draw_process_diagram():
        if not rows:
            draw.text((chart_left, chart_top), "No process data available.", fill="#706960", font=body_font)
            return
        stages = rows[:6]
        box_w = 185
        box_h = 94
        colours = ["#0b0b0b", "#706960", "#f58220", "#2f80ed", "#236b35", "#0b0b0b"]
        positions = [
            (chart_left + 70, chart_top + 35),
            (chart_left + 355, chart_top + 35),
            (chart_left + 640, chart_top + 35),
            (chart_left + 640, chart_top + 245),
            (chart_left + 355, chart_top + 245),
            (chart_left + 70, chart_top + 245),
        ]
        centers = [(x + box_w / 2, y + box_h / 2) for x, y in positions[:len(stages)]]
        for index, (cx, cy) in enumerate(centers):
            nx, ny = centers[(index + 1) % len(centers)]
            draw.line((cx, cy, nx, ny), fill="#d47722", width=4)
            dx = nx - cx
            dy = ny - cy
            length = max(1, (dx * dx + dy * dy) ** .5)
            ux = dx / length
            uy = dy / length
            ax = nx - ux * (box_w / 2 + 12)
            ay = ny - uy * (box_h / 2 + 12)
            draw.polygon([(ax, ay), (ax - ux * 14 - uy * 7, ay - uy * 14 + ux * 7), (ax - ux * 14 + uy * 7, ay - uy * 14 - ux * 7)], fill="#d47722")
        for index, row in enumerate(stages):
            x0, y0 = positions[index]
            fill = colours[index % len(colours)]
            draw.rounded_rectangle((x0, y0, x0 + box_w, y0 + box_h), radius=12, fill=fill, outline="#ffffff", width=2)
            draw.text((x0 + 12, y0 + 10), str(row[0])[:22], fill="#ffffff", font=small_font)
            for line_index, line in enumerate(wrap_text(row[1] if len(row) > 1 else "", 20)[:3]):
                draw.text((x0 + 12, y0 + 38 + line_index * 15), line, fill="#ffffff", font=tiny_font)
        center_x = (chart_left + chart_right) / 2
        center_y = chart_top + 210
        draw.rounded_rectangle((center_x - 190, center_y - 42, center_x + 190, center_y + 42), radius=12, fill="#fbfaf6", outline="#e6ded3", width=2)
        draw.text((center_x - 154, center_y - 24), "Circular methodology loop", fill="#111111", font=small_font)
        draw.text((center_x - 166, center_y - 2), "reviewed skills and aliases feed", fill="#333333", font=tiny_font)
        draw.text((center_x - 152, center_y + 16), "the next analysis refresh", fill="#333333", font=tiny_font)
        draw.text((chart_left, chart_bottom - 28), "The runtime executes in sequence, but the methodology is iterative: review and deployment feed the next run.", fill="#333333", font=small_font)

    def draw_funnel_diagram():
        if not rows:
            draw.text((chart_left, chart_top), "No funnel data available.", fill="#706960", font=body_font)
            return
        stages = rows[:7]
        colours = ["#0b0b0b", "#3c3327", "#7b4b16", "#f58220", "#236b35", "#2f80ed", "#706960"]
        center_x = (chart_left + chart_right) / 2
        max_w = chart_right - chart_left - 180
        min_w = max_w * .35
        stage_h = min(70, (chart_bottom - chart_top - 30) / max(1, len(stages)))
        for index, row in enumerate(stages):
            ratio = index / max(1, len(stages) - 1)
            top_w = max_w - (max_w - min_w) * ratio
            bottom_w = max_w - (max_w - min_w) * min(1, ratio + 1 / max(1, len(stages)))
            y0 = chart_top + index * stage_h
            y1 = y0 + stage_h - 6
            points = [
                (center_x - top_w / 2, y0),
                (center_x + top_w / 2, y0),
                (center_x + bottom_w / 2, y1),
                (center_x - bottom_w / 2, y1),
            ]
            draw.polygon(points, fill=colours[index % len(colours)], outline="#ffffff")
            text_x = center_x - min(top_w, bottom_w) / 2 + 28
            draw.text((text_x, y0 + 10), str(row[0])[:34], fill="#ffffff", font=small_font)
            detail = row[1] if len(row) > 1 else ""
            draw.text((text_x, y0 + 36), short_text(detail, 58), fill="#ffffff", font=tiny_font)
        draw.text((chart_left, chart_bottom - 28), "Funnel direction: broad evidence inputs narrow into validation-ready outputs", fill="#333333", font=small_font)

    def draw_validation_cards():
        if not rows:
            draw.text((chart_left, chart_top), "No validation data available.", fill="#706960", font=body_font)
            return
        cards = rows[:8]
        cols = 2
        card_w = (chart_right - chart_left - 34) / cols
        card_h = min(92, (chart_bottom - chart_top - 32) / max(1, (len(cards) + cols - 1) // cols))
        colours = ["#111111", "#f58220", "#2f80ed", "#236b35"]
        for index, row in enumerate(cards):
            col = index % cols
            r = index // cols
            x0 = chart_left + col * (card_w + 34)
            y0 = chart_top + r * (card_h + 14)
            draw.rounded_rectangle((x0, y0, x0 + card_w, y0 + card_h), radius=10, fill="#fbfaf6", outline="#e6ded3", width=2)
            draw.rectangle((x0, y0, x0 + 8, y0 + card_h), fill=colours[index % len(colours)])
            draw.text((x0 + 18, y0 + 10), str(row[0])[:36], fill="#111111", font=small_font)
            for line_index, line in enumerate(wrap_text(row[1] if len(row) > 1 else "", 44)[:2]):
                draw.text((x0 + 18, y0 + 34 + line_index * 16), line, fill="#333333", font=tiny_font)
            draw.text((x0 + 18, y0 + card_h - 22), f"Output: {short_text(row[2] if len(row) > 2 else '', 48)}", fill="#9a4a08", font=tiny_font)
        legend([("Formula/input", "#111111"), ("Coverage/gap", "#f58220"), ("Correlation/score", "#2f80ed"), ("Review/output", "#236b35")], chart_right - 260, chart_bottom - 110)

    def draw_crisp_dm_loop():
        import math
        stages = rows[:6] or [
            ("Business Understanding", "Curriculum gap question", "Alignment objective"),
            ("Data Understanding", "Courses, modules, jobs", "Evidence audit"),
            ("Data Preparation", "Cleaning and aliases", "Trusted skill evidence"),
            ("Modelling", "NER, embeddings, scoring", "GapResult signals"),
            ("Evaluation", "Human review", "Approved/rejected labels"),
            ("Deployment", "Visuals and report", "Decision support"),
        ]
        center_x = (chart_left + chart_right) / 2
        center_y = chart_top + 160
        radius_x = (chart_right - chart_left) * .30
        radius_y = (chart_bottom - chart_top) * .18
        box_w = 210
        box_h = 78
        colours = ["#0b0b0b", "#706960", "#f58220", "#2f80ed", "#236b35", "#7b4b16"]
        positions = []
        for index in range(len(stages)):
            angle = -math.pi / 2 + index * (math.pi * 2 / len(stages))
            positions.append((center_x + math.cos(angle) * radius_x, center_y + math.sin(angle) * radius_y))
        for index, (x1, y1) in enumerate(positions):
            x2, y2 = positions[(index + 1) % len(positions)]
            draw.line((x1, y1, x2, y2), fill="#d47722", width=4)
            angle = math.atan2(y2 - y1, x2 - x1)
            ax = x2 - math.cos(angle) * (box_w / 2 + 10)
            ay = y2 - math.sin(angle) * (box_h / 2 + 10)
            draw.polygon([
                (ax, ay),
                (ax - math.cos(angle - .55) * 14, ay - math.sin(angle - .55) * 14),
                (ax - math.cos(angle + .55) * 14, ay - math.sin(angle + .55) * 14),
            ], fill="#d47722")
        for index, (row, (cx, cy)) in enumerate(zip(stages, positions)):
            x0 = cx - box_w / 2
            y0 = cy - box_h / 2
            draw.rounded_rectangle((x0, y0, x0 + box_w, y0 + box_h), radius=10, fill=colours[index % len(colours)], outline="#ffffff", width=2)
            title_lines = wrap_text(f"{index + 1}. {row[0]}", 21)[:2]
            for line_index, line in enumerate(title_lines):
                draw.text((x0 + 12, y0 + 9 + line_index * 17), line, fill="#ffffff", font=small_font)
            detail_y = y0 + 43 if len(title_lines) > 1 else y0 + 34
            draw.text((x0 + 12, detail_y), short_text(row[2] if len(row) > 2 else "", 30), fill="#fff7ed", font=tiny_font)
        draw.rounded_rectangle((center_x - 150, center_y - 38, center_x + 150, center_y + 38), radius=12, fill="#fbfaf6", outline="#e6ded3", width=2)
        draw.text((center_x - 122, center_y - 24), "Human-cleaned mining loop", fill="#111111", font=small_font)
        draw.text((center_x - 128, center_y + 2), "approved skills + aliases feed", fill="#333333", font=tiny_font)
        draw.text((center_x - 112, center_y + 20), "hourly refreshed analysis", fill="#333333", font=tiny_font)

        legend_top = chart_bottom - 162
        legend_left = chart_left + 6
        legend_col_w = (chart_right - chart_left - 34) / 2
        row_h = 48
        draw.text((legend_left, legend_top - 24), "Labelled CRISP-DM legend", fill="#111111", font=small_font)
        for index, row in enumerate(stages):
            col = index % 2
            r = index // 2
            x0 = legend_left + col * (legend_col_w + 28)
            y0 = legend_top + r * row_h
            colour = colours[index % len(colours)]
            draw.rounded_rectangle((x0, y0, x0 + 28, y0 + 28), radius=6, fill=colour)
            draw.text((x0 + 9, y0 + 7), str(index + 1), fill="#ffffff", font=tiny_font)
            draw.text((x0 + 38, y0), str(row[0])[:34], fill="#111111", font=small_font)
            for line_index, line in enumerate(wrap_text(f"{row[1]} -> {row[2]}", 66)[:2]):
                draw.text((x0 + 38, y0 + 18 + line_index * 14), line, fill="#333333", font=tiny_font)

    def draw_transformer_architecture():
        boxes = rows or [
            ("Raw text", "job/module sentence"),
            ("Subword tokenizer", "WordPiece/BPE units"),
            ("Token embeddings", "vector per subword"),
            ("Position encodings", "sequence order"),
            ("Transformer encoder", "self-attention blocks"),
            ("Contextual states", "one state per token"),
            ("BIO classifier", "O, B-SKILL, I-SKILL"),
            ("Skill spans", "accepted entities"),
        ]
        box_w = 170
        box_h = 78
        gap = 32
        start_x = chart_left
        y_top = chart_top + 68
        colours = ["#0b0b0b", "#706960", "#2f80ed", "#f58220", "#236b35", "#3f6fa9", "#c83a2f", "#111111"]
        for index, row in enumerate(boxes[:8]):
            x0 = start_x + (index % 4) * (box_w + gap)
            y0 = y_top + (index // 4) * (box_h + 92)
            draw.rounded_rectangle((x0, y0, x0 + box_w, y0 + box_h), radius=10, fill=colours[index % len(colours)], outline="#ffffff", width=2)
            draw.text((x0 + 12, y0 + 10), str(row[0])[:24], fill="#ffffff", font=small_font)
            for line_index, line in enumerate(wrap_text(row[1] if len(row) > 1 else "", 22)[:2]):
                draw.text((x0 + 12, y0 + 36 + line_index * 16), line, fill="#ffffff", font=tiny_font)
            if index < min(8, len(boxes)) - 1:
                if index % 4 != 3:
                    x1 = x0 + box_w
                    y1 = y0 + box_h / 2
                    draw.line((x1 + 6, y1, x1 + gap - 6, y1), fill="#f58220", width=4)
                    draw.polygon([(x1 + gap - 6, y1), (x1 + gap - 18, y1 - 7), (x1 + gap - 18, y1 + 7)], fill="#f58220")
                else:
                    x1 = x0 + box_w / 2
                    y1 = y0 + box_h
                    draw.line((x1, y1 + 8, x1, y1 + 52), fill="#f58220", width=4)
                    draw.polygon([(x1, y1 + 52), (x1 - 7, y1 + 40), (x1 + 7, y1 + 40)], fill="#f58220")
        # Self-attention inset
        inset_x = chart_right - 380
        inset_y = chart_bottom - 160
        draw.rounded_rectangle((inset_x, inset_y, inset_x + 330, inset_y + 110), radius=10, fill="#fbfaf6", outline="#e6ded3", width=2)
        draw.text((inset_x + 14, inset_y + 10), "Self-attention core", fill="#111111", font=small_font)
        tokens = ["x1", "x2", "x3", "x4"]
        for i, token in enumerate(tokens):
            px = inset_x + 58 + i * 70
            py = inset_y + 64
            draw.ellipse((px - 15, py - 15, px + 15, py + 15), fill="#2f80ed", outline="#ffffff", width=2)
            draw.text((px - 7, py - 7), token, fill="#ffffff", font=tiny_font)
            for j in range(i + 1, len(tokens)):
                qx = inset_x + 58 + j * 70
                draw.line((px, py, qx, py), fill="#d47722", width=1)
        draw.text((chart_left, chart_bottom - 28), "Transformer flow: tokens become contextual states, then BIO labels produce skill spans.", fill="#333333", font=small_font)

    def draw_event_sequence():
        events = rows or []
        if not events:
            draw.text((chart_left, chart_top), "No event flow data available.", fill="#706960", font=body_font)
            return
        actors = []
        for row in events:
            if len(row) >= 3:
                if row[0] not in actors:
                    actors.append(row[0])
                if row[1] not in actors:
                    actors.append(row[1])
        actors = actors[:10]
        actor_gap = (chart_right - chart_left - 120) / max(1, len(actors) - 1)
        actor_half_width = 58 if len(actors) <= 7 else 48
        x_positions = {actor: chart_left + 60 + idx * actor_gap for idx, actor in enumerate(actors)}
        for actor, x_pos in x_positions.items():
            draw.rounded_rectangle((x_pos - actor_half_width, chart_top, x_pos + actor_half_width, chart_top + 42), radius=8, fill="#111111", outline="#ffffff")
            draw.text((x_pos - actor_half_width + 8, chart_top + 13), short_text(actor, 14), fill="#ffffff", font=tiny_font)
            draw.line((x_pos, chart_top + 48, x_pos, chart_bottom - 32), fill="#d8d2c8", width=2)
        y_step = max(26, (chart_bottom - chart_top - 105) / max(1, len(events)))
        for index, row in enumerate(events[:18]):
            if len(row) < 3:
                continue
            source, target, message = row[:3]
            if source not in x_positions or target not in x_positions:
                continue
            y_pos = chart_top + 76 + index * y_step
            x1 = x_positions[source]
            x2 = x_positions[target]
            draw.line((x1, y_pos, x2, y_pos), fill="#f58220", width=3)
            arrow = 8 if x2 >= x1 else -8
            draw.polygon([(x2, y_pos), (x2 - arrow, y_pos - 6), (x2 - arrow, y_pos + 6)], fill="#f58220")
            draw.text((min(x1, x2) + 8, y_pos - 18), short_text(message, 46), fill="#333333", font=tiny_font)
        draw.text((chart_left, chart_bottom - 24), "Event order follows analysis/services.py: load data, score, save results, then export visuals.", fill="#333333", font=small_font)

    if "skill demand vs curriculum" in title_lower:
        draw_line()
    elif "methodology diagram 1" in title_lower:
        draw_process_diagram()
    elif "methodology diagram 2" in title_lower:
        draw_funnel_diagram()
    elif "methodology diagram 3" in title_lower:
        draw_validation_cards()
    elif "crisp-dm" in title_lower or "data-mining process" in title_lower or "data mining process" in title_lower:
        draw_crisp_dm_loop()
    elif "transformer architecture" in title_lower:
        draw_transformer_architecture()
    elif "flow of events" in title_lower or "runtime event" in title_lower:
        draw_event_sequence()
    elif "divergence" in title_lower:
        draw_diverging_bars()
    elif "network" in title_lower or "cluster" in title_lower or "linked to type" in title_lower:
        draw_network()
    elif "heatmap" in title_lower or "cross-tab" in title_lower or "density" in title_lower:
        draw_heatmap()
    elif "scatter" in title_lower or "matched vs missing" in title_lower:
        draw_scatter()
    elif "forecast" in title_lower or "trend" in title_lower:
        if rows and len(rows[0]) >= 3:
            rows = [(row[1], row[2]) for row in rows]
        draw_line()
    elif "skill type and source" in title_lower:
        if rows and len(rows[0]) >= 3:
            rows = [(f"{row[0]}: {row[1]}", row[2]) for row in rows]
        draw_bars(grouped=False)
    elif fallback_headers and len(fallback_headers) >= 3 and any("course" in str(header).lower() or "job" in str(header).lower() for header in fallback_headers[1:]):
        draw_bars(grouped=True)
    elif fallback_headers and len(fallback_headers) >= 3:
        draw_bars(grouped=True)
    else:
        draw_bars(grouped=False)

    output = BytesIO()
    image.save(output, format="PNG")
    output.seek(0)
    return output


def figure_reference_text(title, fallback_headers=None, fallback_rows=None):
    headers = list(fallback_headers or [])
    rows = [tuple(row) for row in (fallback_rows or []) if row]
    title_lower = title.lower()
    if not rows:
        return f"Reference to {title}: this figure has no source rows in the current export, so it should be interpreted as an empty visual for this run."

    def number(value, default=0):
        try:
            return float(str(value).replace("%", "").replace(",", "").strip())
        except (TypeError, ValueError):
            return default

    def label(value):
        return str(value or "unlabelled")

    if "heatmap" in title_lower or "cross-tab" in title_lower or "density" in title_lower:
        strongest = max(rows, key=lambda row: number(row[2] if len(row) > 2 else 0))
        x_name = headers[1] if len(headers) > 1 else "column"
        y_name = headers[0] if headers else "row"
        value_name = headers[2] if len(headers) > 2 else "value"
        return (
            f"Reference to {title}: read the y-axis as {y_name} and the x-axis as {x_name}; darker cells indicate higher {value_name}. "
            f"The strongest visible cell is {label(strongest[0])} by {label(strongest[1] if len(strongest) > 1 else '')}, with {value_name} {strongest[2] if len(strongest) > 2 else '0'}."
        )
    if "scatter" in title_lower or "matched vs missing" in title_lower:
        highest_missing = max(rows, key=lambda row: number(row[2] if len(row) > 2 else 0))
        return (
            f"Reference to {title}: the x-axis shows {headers[1] if len(headers) > 1 else 'matched evidence'} and the y-axis shows {headers[2] if len(headers) > 2 else 'missing evidence'}. "
            f"The highest missing-evidence point is {label(highest_missing[0])}, with x={highest_missing[1] if len(highest_missing) > 1 else 0} and y={highest_missing[2] if len(highest_missing) > 2 else 0}."
        )
    if "network" in title_lower or "cluster" in title_lower or "linked to type" in title_lower:
        return (
            f"Reference to {title}: nodes represent the labels in the source rows and lines represent relationships between them. "
            f"The figure is built from {len(rows)} relationship or node rows; dense areas show where skills, roles, sources, or courses are repeatedly connected."
        )
    if "forecast" in title_lower or "trend" in title_lower:
        latest = rows[-1]
        return (
            f"Reference to {title}: the x-axis follows the dated evidence sequence and the y-axis shows record counts. "
            f"The final plotted row is {label(latest[0])}, with value {latest[-1] if latest else 0}."
        )
    if len(headers) >= 3:
        first_value = max(rows, key=lambda row: number(row[1] if len(row) > 1 else 0))
        second_value = max(rows, key=lambda row: number(row[2] if len(row) > 2 else 0))
        return (
            f"Reference to {title}: the x-axis lists {headers[0]}, while the y-axis shows the numeric evidence values. "
            f"The largest {headers[1]} value is {label(first_value[0])} ({first_value[1]}), and the largest {headers[2]} value is {label(second_value[0])} ({second_value[2]})."
        )
    strongest = max(rows, key=lambda row: number(row[1] if len(row) > 1 else 0))
    return (
        f"Reference to {title}: the x-axis lists {headers[0] if headers else 'categories'} and the y-axis shows {headers[1] if len(headers) > 1 else 'counts'}. "
        f"The largest plotted value is {label(strongest[0])}, with {strongest[1] if len(strongest) > 1 else 0}."
    )


def add_figure_caption(document, title, caption, fallback_headers=None, fallback_rows=None):
    document.add_paragraph(caption)
    paragraph = document.add_paragraph()
    paragraph.add_run(figure_reference_text(title, fallback_headers, fallback_rows)).italic = True


def add_figure(
    document,
    title,
    caption,
    fig,
    fallback_headers=None,
    fallback_rows=None,
    width=6.4,
    asset_dir=None,
    image_name=None,
    manifest_rows=None,
):
    from docx.shared import Inches

    document.add_heading(title, level=2)
    image = plotly_to_image(fig)
    if image:
        if asset_dir and image_name:
            image_path = asset_dir / f"{safe_filename(image_name)}.png"
            image_path.write_bytes(image.getvalue())
            document.add_picture(str(image_path), width=Inches(width))
            if manifest_rows is not None:
                manifest_rows.append({
                    "title": title,
                    "file": f"images/{image_path.name}",
                    "caption": caption,
                })
        else:
            document.add_picture(image, width=Inches(width))
        add_figure_caption(document, title, caption, fallback_headers, fallback_rows)
        return True
    fallback_image = fallback_png(title, caption, fallback_headers, fallback_rows)
    if fallback_image and asset_dir and image_name:
        image_path = asset_dir / f"{safe_filename(image_name)}.png"
        image_path.write_bytes(fallback_image.getvalue())
        document.add_picture(str(image_path), width=Inches(width))
        if manifest_rows is not None:
            manifest_rows.append({
                "title": title,
                "file": f"images/{image_path.name}",
                "caption": f"{caption} Rendered as a cached PNG visual from system data.",
            })
        add_figure_caption(document, title, caption, fallback_headers, fallback_rows)
        return True
    document.add_paragraph(f"{caption} Image rendering was unavailable, so the source table is shown below.")
    document.add_paragraph(figure_reference_text(title, fallback_headers, fallback_rows))
    if fallback_headers and fallback_rows:
        add_simple_table(document, fallback_headers, fallback_rows)
    return False


def create_bar_figure(title, labels, values, x_title="", y_title="Count", color="#f58220"):
    go = plotly_go()
    if go is None:
        return None

    fig = go.Figure(go.Bar(
        x=labels,
        y=values,
        marker={"color": color},
        text=values,
        textposition="outside",
    ))
    fig.update_layout(
        title={"text": title, "x": 0.02, "xanchor": "left"},
        margin={"l": 60, "r": 30, "t": 70, "b": 130},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": x_title, "tickangle": -35},
        yaxis={"title": y_title, "gridcolor": "#f2eee8"},
        showlegend=False,
    )
    return fig


def create_horizontal_bar_figure(title, rows, x_title="Records", color="#f58220", limit=15):
    go = plotly_go()
    if go is None:
        return None
    selected = list(rows or [])[:limit]
    labels = [row[0] for row in selected][::-1]
    values = [row[1] for row in selected][::-1]
    fig = go.Figure(go.Bar(
        x=values,
        y=labels,
        orientation="h",
        marker={"color": color},
        text=values,
        textposition="outside",
    ))
    fig.update_layout(
        title={"text": title, "x": 0.02, "xanchor": "left"},
        margin={"l": 210, "r": 40, "t": 70, "b": 60},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": x_title, "gridcolor": "#f2eee8"},
        yaxis={"automargin": True},
        showlegend=False,
    )
    return fig


def dashboard_skill_line_rows(run=None):
    job_counter = Counter()
    course_counter = Counter()
    if run:
        for row in SkillMatrix.objects.filter(run=run, source="jobs"):
            job_counter[row.skill] += row.frequency
        for row in SkillMatrix.objects.filter(run=run, source="courses"):
            course_counter[row.skill] += row.frequency
    for job in JobAdvert.objects.only("skills_extracted", "skill_entities"):
        job_counter.update(skill for skill in entity_skill_names(job.skill_entities or job.skills_extracted) if skill)
    for module in Module.objects.only("skills_extracted", "skill_entities"):
        course_counter.update(skill for skill in entity_skill_names(module.skill_entities or module.skills_extracted) if skill)
    labels = []
    for skill, _count in job_counter.most_common(10):
        if skill not in labels:
            labels.append(skill)
    for skill, _count in course_counter.most_common(10):
        if skill not in labels:
            labels.append(skill)
    labels = labels[:10]
    return [(label, job_counter.get(label, 0), course_counter.get(label, 0)) for label in labels]


def entity_skill_names(raw_entities):
    names = []
    for raw in raw_entities or []:
        if isinstance(raw, dict):
            if raw.get("skill_type") == "exclude" or raw.get("label") == "exclude" or raw.get("label_status") == "candidate":
                continue
            skill = raw.get("skill") or raw.get("text") or ""
        else:
            skill = str(raw or "")
        skill = " ".join(skill.lower().replace("-", " ").split())
        if skill:
            names.append(skill)
    return names


def dashboard_skill_line_figure(run=None):
    go = plotly_go()
    if go is None:
        return None
    rows = dashboard_skill_line_rows(run)
    labels = [row[0] for row in rows]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        name="Jobs",
        x=labels,
        y=[row[1] for row in rows],
        mode="lines+markers",
        line={"color": "#0b0b0b", "width": 3, "shape": "spline"},
        marker={"size": 8},
        fill="tozeroy",
        fillcolor="rgba(11,11,11,.08)",
    ))
    fig.add_trace(go.Scatter(
        name="Courses",
        x=labels,
        y=[row[2] for row in rows],
        mode="lines+markers",
        line={"color": "#f58220", "width": 3, "shape": "spline"},
        marker={"size": 8},
        fill="tozeroy",
        fillcolor="rgba(245,130,32,.18)",
    ))
    fig.update_layout(
        title={"text": "Skill Demand vs Curriculum", "x": 0.02, "xanchor": "left"},
        margin={"l": 60, "r": 30, "t": 70, "b": 140},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"tickangle": -35},
        yaxis={"title": "Evidence records", "rangemode": "tozero", "gridcolor": "#f2eee8"},
        legend={"orientation": "h", "y": 1.08},
    )
    return fig


def score_band_heatmap_figure(visual_data):
    go = plotly_go()
    if go is None:
        return None
    rows = visual_data.get("heatmap_rows", [])
    bands = visual_data.get("score_bands", [])
    z = [[cell["count"] for cell in row.get("cells", [])] for row in rows]
    labels = [[str(cell["count"]) for cell in row.get("cells", [])] for row in rows]
    y_labels = [row["course"].code or row["course"].name for row in rows]
    fig = go.Figure(go.Heatmap(
        x=[f"{band}%" for band in bands],
        y=y_labels,
        z=z,
        text=labels,
        texttemplate="%{text}",
        colorscale=[[0, "#f1f3f5"], [0.35, "#ffd2aa"], [0.7, "#f58220"], [1, "#974800"]],
        colorbar={"title": "Jobs"},
        xgap=3,
        ygap=3,
    ))
    fig.update_layout(
        title={"text": "Course Alignment Score-Band Heatmap", "x": 0.02, "xanchor": "left"},
        margin={"l": 190, "r": 40, "t": 70, "b": 80},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": "Score band"},
        yaxis={"autorange": "reversed", "automargin": True},
        height=max(520, len(rows) * 42 + 140),
    )
    return fig


def score_distribution_figure(results):
    rows = score_distribution_rows(results)
    return create_bar_figure("Course-to-Job Score Distribution", [row[0] for row in rows], [row[1] for row in rows], "Score band", "Comparisons", "#2f80ed")


def score_distribution_rows(results):
    labels = ["0-20", "20-40", "40-60", "60-80", "80-100"]
    counts = dict.fromkeys(labels, 0)
    for result in results:
        score = result.similarity_percent
        if score < 20:
            counts["0-20"] += 1
        elif score < 40:
            counts["20-40"] += 1
        elif score < 60:
            counts["40-60"] += 1
        elif score < 80:
            counts["60-80"] += 1
        else:
            counts["80-100"] += 1
    return [(label, counts[label]) for label in labels]


def top_skill_rows(matrix_rows, limit=15):
    return [(row.skill, row.frequency) for row in matrix_rows[:limit]]


def skill_gap_rows(visual_data, limit=15):
    rows = visual_data.get("skill_suggestion_matrix", [])[:limit]
    return [
        (
            row["skill"],
            row["demand_count"],
            row["matched_count"],
            row["missing_count"],
            f'{row["gap_percent"]}%',
        )
        for row in rows
    ]


def correlation_rows(visual_data, limit=12):
    matrix = visual_data.get("skill_correlation_matrix", {})
    skills = matrix.get("skills", [])
    values = matrix.get("values", [])
    rows = []
    for row_index, row_skill in enumerate(skills):
        for col_index, col_skill in enumerate(skills):
            if row_index >= len(values) or col_index >= len(values[row_index]):
                continue
            rows.append((row_skill, col_skill, values[row_index][col_index]))
            if len(rows) >= limit:
                return rows
    return rows


def role_divergence_rows(visual_data, limit=18):
    rows = []
    for role in (visual_data.get("role_skill_divergence", {}).get("roles") or []):
        for skill, value in zip(role.get("skills", []), role.get("values", [])):
            rows.append((role["role"], skill, value))
            if len(rows) >= limit:
                return rows
    return rows


def skill_gap_figure(visual_data):
    go = plotly_go()
    if go is None:
        return None

    rows = visual_data.get("skill_suggestion_matrix", [])[:15]
    labels = [row["skill"] for row in rows]
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Covered", x=labels, y=[row["matched_count"] for row in rows], marker={"color": "#236b35"}))
    fig.add_trace(go.Bar(name="Missing", x=labels, y=[row["missing_count"] for row in rows], marker={"color": "#f58220"}))
    fig.update_layout(
        title={"text": "Skill Coverage and Gap Evidence", "x": 0.02, "xanchor": "left"},
        barmode="group",
        margin={"l": 60, "r": 30, "t": 70, "b": 140},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"tickangle": -35},
        yaxis={"title": "Evidence count", "gridcolor": "#f2eee8"},
        legend={"orientation": "h", "y": 1.08},
    )
    return fig


def course_job_crosstab_rows(results, course_limit=18, job_limit=24):
    selected = sorted(results, key=lambda item: item.similarity_score, reverse=True)
    course_ids = []
    job_ids = []
    for result in selected:
        if result.course_id not in course_ids:
            course_ids.append(result.course_id)
        if result.job_id not in job_ids:
            job_ids.append(result.job_id)
        if len(course_ids) >= course_limit and len(job_ids) >= job_limit:
            break
    course_ids = course_ids[:course_limit]
    job_ids = job_ids[:job_limit]
    courses = []
    jobs = []
    seen = set()
    for result in selected:
        if result.course_id in course_ids and result.course_id not in seen:
            courses.append(result.course)
            seen.add(result.course_id)
    seen = set()
    for result in selected:
        if result.job_id in job_ids and result.job_id not in seen:
            jobs.append(result.job)
            seen.add(result.job_id)
    score_map = {(result.course_id, result.job_id): result.similarity_percent for result in results}
    return courses, jobs, score_map


def course_job_crosstab_figure(results):
    go = plotly_go()
    if go is None:
        return None
    courses, jobs, score_map = course_job_crosstab_rows(results)
    z = [[score_map.get((course.id, job.id)) for job in jobs] for course in courses]
    fig = go.Figure(go.Heatmap(
        x=[job.title[:32] for job in jobs],
        y=[course.code or course.name[:32] for course in courses],
        z=z,
        colorscale=[[0, "#f4f1eb"], [0.35, "#ffd5ad"], [0.65, "#f58220"], [1, "#111111"]],
        zmin=0,
        zmax=100,
        colorbar={"title": "Score %"},
        xgap=2,
        ygap=2,
        hovertemplate="Course: %{y}<br>Job: %{x}<br>Score: %{z:.1f}%<extra></extra>",
    ))
    fig.update_layout(
        title={"text": "Course-to-Job Similarity Cross-tab", "x": 0.02, "xanchor": "left"},
        margin={"l": 170, "r": 34, "t": 70, "b": 170},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"tickangle": -38, "automargin": True},
        yaxis={"automargin": True, "autorange": "reversed"},
        height=max(520, len(courses) * 34 + 170),
    )
    return fig


def similarity_network_figure(results):
    go = plotly_go()
    if go is None:
        return None
    try:
        import networkx as nx
    except Exception:
        nx = None
    selected = sorted(results, key=lambda item: item.similarity_score, reverse=True)[:40]
    graph_nodes = {}
    edges = []
    for result in selected:
        course_id = f"course-{result.course_id}"
        job_id = f"job-{result.job_id}"
        graph_nodes[course_id] = {"label": result.course.code or result.course.name[:28], "group": "Course"}
        graph_nodes[job_id] = {"label": result.job.title[:28], "group": "Job advert"}
        edges.append((course_id, job_id, result.similarity_percent))
    if nx:
        graph = nx.Graph()
        graph.add_nodes_from(graph_nodes)
        for left, right, value in edges:
            graph.add_edge(left, right, value=value)
        positions = nx.spring_layout(graph, seed=42, k=0.85) if graph.number_of_nodes() else {}
    else:
        positions = {}
        course_index = 0
        job_index = 0
        for node_id, attrs in graph_nodes.items():
            if attrs["group"] == "Course":
                positions[node_id] = (-1, course_index)
                course_index += 1
            else:
                positions[node_id] = (1, job_index)
                job_index += 1

    edge_x = []
    edge_y = []
    for left, right, _value in edges:
        left_pos = positions.get(left, (0, 0))
        right_pos = positions.get(right, (0, 0))
        edge_x.extend([left_pos[0], right_pos[0], None])
        edge_y.extend([left_pos[1], right_pos[1], None])

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=edge_x,
        y=edge_y,
        mode="lines",
        line={"color": "rgba(212,119,34,.45)", "width": 1.4},
        hoverinfo="skip",
        showlegend=False,
    ))
    for group, color in [("Course", "#111111"), ("Job advert", "#f58220")]:
        ids = [node_id for node_id, attrs in graph_nodes.items() if attrs["group"] == group]
        fig.add_trace(go.Scatter(
            name=group,
            x=[positions.get(node_id, (0, 0))[0] for node_id in ids],
            y=[positions.get(node_id, (0, 0))[1] for node_id in ids],
            mode="markers+text",
            text=[graph_nodes[node_id]["label"] for node_id in ids],
            textposition="top center",
            marker={"size": 15, "color": color, "line": {"color": "#ffffff", "width": 1}},
            hovertemplate="%{text}<extra></extra>",
        ))
    fig.update_layout(
        title={"text": "Cosine Similarity Network", "x": 0.02, "xanchor": "left"},
        margin={"l": 40, "r": 40, "t": 70, "b": 40},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"visible": False},
        yaxis={"visible": False},
        legend={"orientation": "h", "y": 1.05},
        height=620,
    )
    return fig


def scatter_figure(visual_data):
    go = plotly_go()
    if go is None:
        return None

    points = visual_data.get("scatter_points", [])
    fig = go.Figure(go.Scatter(
        x=[point["x"] for point in points],
        y=[point["y"] for point in points],
        mode="markers+text",
        text=[point["label"] for point in points],
        textposition="top center",
        marker={
            "size": [max(10, min(38, point["score"] / 2)) for point in points],
            "color": [point["score"] for point in points],
            "colorscale": [[0, "#f1e7dd"], [0.5, "#f58220"], [1, "#236b35"]],
            "showscale": True,
            "colorbar": {"title": "Avg score"},
        },
    ))
    fig.update_layout(
        title={"text": "Matched vs Missing Skill Evidence", "x": 0.02, "xanchor": "left"},
        margin={"l": 70, "r": 40, "t": 70, "b": 70},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": "Matched evidence", "gridcolor": "#f2eee8"},
        yaxis={"title": "Missing evidence", "gridcolor": "#f2eee8"},
    )
    return fig


def heatmap_figure(visual_data):
    go = plotly_go()
    if go is None:
        return None

    matrix = visual_data.get("skill_correlation_matrix", {})
    skills = matrix.get("skills", [])
    z = matrix.get("values", [])
    fig = go.Figure(go.Heatmap(
        z=z,
        x=skills,
        y=skills,
        text=matrix.get("labels", []),
        texttemplate="%{text}",
        colorscale=[[0, "#3f6fa9"], [0.5, "#edf1f5"], [1, "#b92f24"]],
        colorbar={"title": "Pearson r"},
        zmin=-1,
        zmax=1,
    ))
    fig.update_layout(
        title={"text": "Correlation Structure Between Skills", "x": 0.02, "xanchor": "left"},
        margin={"l": 160, "r": 30, "t": 70, "b": 150},
        paper_bgcolor="white",
        xaxis={"tickangle": -40},
        yaxis={"autorange": "reversed"},
    )
    return fig


def role_divergence_figure(visual_data):
    go = plotly_go()
    if go is None:
        return None

    roles = (visual_data.get("role_skill_divergence", {}).get("roles") or [])[:4]
    fig = go.Figure()
    y_labels = []
    x_values = []
    colours = []
    for role in roles:
        for skill, value in zip(role.get("skills", [])[:6], role.get("values", [])[:6]):
            y_labels.append(f"{role['role']} | {skill}")
            x_values.append(value)
            colours.append("#c83a2f" if value >= 0 else "#3f6fa9")
    fig.add_trace(go.Bar(
        x=x_values,
        y=y_labels,
        orientation="h",
        marker={"color": colours},
        text=[f"{value:.2f}" for value in x_values],
        textposition="outside",
    ))
    max_abs = max([abs(value) for value in x_values] or [0.1])
    fig.update_layout(
        title={"text": "Role Skill Divergence", "x": 0.02, "xanchor": "left"},
        margin={"l": 260, "r": 60, "t": 70, "b": 60},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": "Skill share difference from baseline", "range": [-max_abs * 1.35, max_abs * 1.35], "zeroline": True, "zerolinewidth": 2, "gridcolor": "#f2eee8"},
        yaxis={"automargin": True},
        showlegend=False,
        height=max(520, len(y_labels) * 28 + 120),
    )
    return fig


def single_role_divergence_figure(role):
    go = plotly_go()
    if go is None or not role:
        return None
    values = role.get("values", [])
    skills = role.get("skills", [])
    max_abs = max([abs(value) for value in values] or [0.1])
    fig = go.Figure(go.Bar(
        x=values,
        y=skills,
        orientation="h",
        marker={"color": ["#c83a2f" if value >= 0 else "#3f6fa9" for value in values]},
        text=[f"{value:.2f}" for value in values],
        textposition="outside",
    ))
    fig.update_layout(
        title={"text": role.get("role", "Role profile"), "x": 0.5, "xanchor": "center"},
        margin={"l": 210, "r": 110, "t": 70, "b": 50},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"range": [-max_abs * 1.35, max_abs * 1.35], "zeroline": True, "zerolinewidth": 2, "gridcolor": "#f2eee8"},
        yaxis={"automargin": True},
        showlegend=False,
        height=max(360, len(skills) * 38 + 120),
    )
    return fig


def report_skill_entity_rows(run=None):
    rows = []
    for job in JobAdvert.objects.order_by("title", "id"):
        label = f"{job.title} @ {job.company}" if job.company else job.title
        for skill in entity_skill_names(job.skill_entities or job.skills_extracted):
            rows.append({
                "skill": skill,
                "skill_type": "job",
                "source_type": "job",
                "sector": job.category or "Unclassified",
                "extracted_year": job.date_posted.year if job.date_posted else (job.created_at.year if job.created_at else None),
                "source_label": label,
            })
    for module in Module.objects.select_related("course").order_by("course__code", "order", "name", "id"):
        for skill in entity_skill_names(module.skill_entities or module.skills_extracted):
            rows.append({
                "skill": skill,
                "skill_type": "course",
                "source_type": "module",
                "sector": module.course.university_name or "Course module",
                "extracted_year": module.created_at.year if getattr(module, "created_at", None) else None,
                "source_label": module.name,
            })
    if not rows and run:
        for matrix in SkillMatrix.objects.filter(run=run).order_by("source", "-frequency", "skill"):
            rows.extend({
                "skill": matrix.skill,
                "skill_type": matrix.source[:-1] if matrix.source.endswith("s") else matrix.source,
                "source_type": "job" if matrix.source == "jobs" else "module",
                "sector": "Analysis aggregate",
                "extracted_year": run.created_at.year if run.created_at else None,
                "source_label": run.name,
            } for _index in range(max(1, matrix.frequency)))
    return rows


def data_export_top_skills_figure(rows):
    return create_horizontal_bar_figure(
        "Top Skill Evidence",
        Counter(row["skill"] for row in rows).most_common(15),
        "Records",
        "#f58220",
    )


def data_export_type_source_figure(rows):
    go = plotly_go()
    if go is None:
        return None
    type_pairs = Counter(row["skill_type"] for row in rows).most_common(10)
    source_pairs = Counter(row["source_type"] for row in rows).most_common(10)
    fig = go.Figure()
    fig.add_trace(go.Pie(
        labels=[row[0] for row in type_pairs],
        values=[row[1] for row in type_pairs],
        hole=.48,
        marker={"colors": ["#111111", "#f58220", "#2f80ed", "#6f6a60", "#9e2f18"]},
        domain={"x": [0, .48], "y": [0, 1]},
        name="Skill type",
    ))
    fig.add_trace(go.Bar(
        x=[row[0] for row in source_pairs],
        y=[row[1] for row in source_pairs],
        marker={"color": "#2f80ed"},
        xaxis="x2",
        yaxis="y2",
        name="Source",
    ))
    fig.update_layout(
        title={"text": "Skill Type and Source", "x": 0.02, "xanchor": "left"},
        margin={"l": 50, "r": 30, "t": 70, "b": 90},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis2={"domain": [.58, 1], "anchor": "y2", "tickangle": -20},
        yaxis2={"domain": [0, 1], "anchor": "x2", "gridcolor": "#f2eee8"},
        showlegend=False,
    )
    return fig


def data_export_skill_heatmap_figure(rows):
    go = plotly_go()
    if go is None:
        return None
    sources = sorted({row["source_type"] or "unknown" for row in rows})
    types = sorted({row["skill_type"] or "unknown" for row in rows})
    counts = Counter((row["source_type"] or "unknown", row["skill_type"] or "unknown") for row in rows)
    fig = go.Figure(go.Heatmap(
        x=sources,
        y=types,
        z=[[counts.get((source, skill_type), 0) for source in sources] for skill_type in types],
        colorscale=[[0, "#fff7ed"], [.35, "#fdba74"], [.7, "#f58220"], [1, "#111111"]],
        colorbar={"title": "Records"},
        hovertemplate="Source: %{x}<br>Type: %{y}<br>Records: %{z}<extra></extra>",
    ))
    fig.update_layout(
        title={"text": "NER Skill Evidence Density", "x": 0.02, "xanchor": "left"},
        margin={"l": 110, "r": 30, "t": 70, "b": 80},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": "Source"},
        yaxis={"title": "Skill type", "automargin": True},
    )
    return fig


def data_export_forecast_figure(rows):
    go = plotly_go()
    if go is None:
        return None
    dated = [row for row in rows if row.get("extracted_year")]
    if not dated:
        return None
    skills = [skill for skill, _count in Counter(row["skill"] for row in dated).most_common(6)]
    years = sorted({row["extracted_year"] for row in dated})
    next_year = max(years) + 1
    fig = go.Figure()
    colors = ["#111111", "#f58220", "#2f80ed", "#236b35", "#9e2f18", "#706960"]
    for index, skill in enumerate(skills):
        counts = [sum(1 for row in dated if row["skill"] == skill and row["extracted_year"] == year) for year in years]
        slope = (counts[-1] - counts[0]) / max(1, len(counts) - 1) if len(counts) > 1 else max(1, counts[0] * .15)
        forecast = max(0, round(counts[-1] + slope))
        color = colors[index % len(colors)]
        fig.add_trace(go.Scatter(name=skill, x=years, y=counts, mode="lines+markers", line={"color": color, "width": 3}))
        fig.add_trace(go.Scatter(name=f"{skill} forecast", x=[years[-1], next_year], y=[counts[-1], forecast], mode="lines+markers", line={"color": color, "width": 2, "dash": "dot"}, showlegend=False))
    fig.update_layout(
        title={"text": "Skill Evidence Trend and One-Year Forecast", "x": 0.02, "xanchor": "left"},
        margin={"l": 60, "r": 30, "t": 70, "b": 100},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"title": "Extraction year", "dtick": 1, "gridcolor": "#f2eee8"},
        yaxis={"title": "Skill records", "rangemode": "tozero", "gridcolor": "#f2eee8"},
        legend={"orientation": "h", "y": -0.2},
    )
    return fig


def data_export_skill_network_figure(rows):
    go = plotly_go()
    if go is None:
        return None
    top_skills = {skill for skill, _count in Counter(row["skill"] for row in rows).most_common(18)}
    selected = [row for row in rows if row["skill"] in top_skills][:140]
    if not selected:
        return None
    skills = sorted({row["skill"] for row in selected})
    types = sorted({row["skill_type"] or "unknown" for row in selected})
    sources = sorted({row["source_type"] or "unknown" for row in selected})
    nodes = []
    skill_step = max(1.1, 18 / max(1, len(skills) - 1))
    hub_step = 3.6
    for index, label in enumerate(skills):
        nodes.append({
            "id": f"skill:{label}",
            "label": label,
            "kind": "skill",
            "count": sum(1 for row in selected if row["skill"] == label),
            "x": (index - (len(skills) - 1) / 2) * skill_step,
            "y": 2 + (index % 2) * .18,
        })
    for index, label in enumerate(types):
        nodes.append({
            "id": f"type:{label}",
            "label": label,
            "kind": "type",
            "count": sum(1 for row in selected if (row["skill_type"] or "unknown") == label),
            "x": (index - (len(types) - 1) / 2) * hub_step,
            "y": .95,
        })
    for index, label in enumerate(sources):
        nodes.append({
            "id": f"source:{label}",
            "label": label,
            "kind": "source",
            "count": sum(1 for row in selected if (row["source_type"] or "unknown") == label),
            "x": (index - (len(sources) - 1) / 2) * hub_step,
            "y": -.1,
        })
    node_map = {node["id"]: node for node in nodes}
    edge_x = []
    edge_y = []
    seen_edges = set()
    for row in selected:
        for target in (f"type:{row['skill_type'] or 'unknown'}", f"source:{row['source_type'] or 'unknown'}"):
            key = (f"skill:{row['skill']}", target)
            if key in seen_edges:
                continue
            seen_edges.add(key)
            left = node_map.get(key[0])
            right = node_map.get(key[1])
            if left and right:
                edge_x.extend([left["x"], right["x"], None])
                edge_y.extend([left["y"], right["y"], None])
    color_by_kind = {"skill": "#f58220", "type": "#111111", "source": "#2f80ed"}
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=edge_x, y=edge_y, mode="lines", line={"color": "rgba(112,105,96,.24)", "width": 1}, hoverinfo="skip", showlegend=False))
    for kind in ("skill", "type", "source"):
        kind_nodes = [node for node in nodes if node["kind"] == kind]
        fig.add_trace(go.Scatter(
            name="Skill nodes" if kind == "skill" else f"{kind.title()} hubs",
            x=[node["x"] for node in kind_nodes],
            y=[node["y"] for node in kind_nodes],
            text=[node["label"] for node in kind_nodes],
            mode="markers+text",
            textposition="bottom center",
            marker={
                "size": [min(26, 11 + node["count"] * 1.5) if kind == "skill" else 20 for node in kind_nodes],
                "color": color_by_kind[kind],
                "opacity": .9,
                "line": {"color": "#ffffff", "width": 1},
            },
            hovertemplate="%{text}<extra></extra>",
        ))
    fig.update_layout(
        title={"text": "Skills Linked to Type and Source", "x": 0.02, "xanchor": "left"},
        margin={"l": 30, "r": 30, "t": 70, "b": 80},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"visible": False},
        yaxis={"visible": False},
        legend={"orientation": "h", "y": -0.12, "x": .5, "xanchor": "center"},
        height=560,
    )
    return fig


def data_export_cluster_figure(rows):
    go = plotly_go()
    if go is None:
        return None
    pairs = Counter(row["skill"] for row in rows).most_common(36)
    if not pairs:
        return None
    skills = [skill for skill, _count in pairs]
    nodes = []
    for index, (skill, count) in enumerate(pairs):
        angle = index * 0.72
        radius = 1 + (index % 6) * .18
        skill_rows = [row for row in rows if row["skill"] == skill]
        skill_type = Counter(row["skill_type"] for row in skill_rows).most_common(1)[0][0]
        nodes.append({
            "skill": skill,
            "count": count,
            "skill_type": skill_type,
            "tokens": set(skill.split()),
            "x": radius * __import__("math").cos(angle),
            "y": radius * __import__("math").sin(angle),
        })
    edge_x = []
    edge_y = []
    for left_index, left in enumerate(nodes):
        for right_index in range(left_index + 1, len(nodes)):
            right = nodes[right_index]
            token_score = len(left["tokens"] & right["tokens"]) / max(1, len(left["tokens"] | right["tokens"]))
            type_score = .18 if left["skill_type"] == right["skill_type"] else 0
            if token_score + type_score < .28:
                continue
            edge_x.extend([left["x"], right["x"], None])
            edge_y.extend([left["y"], right["y"], None])
    color_by_type = {"technical": "#2f80ed", "business": "#f58220", "soft": "#236b35", "course": "#236b35", "job": "#f58220"}
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=edge_x, y=edge_y, mode="lines", line={"color": "rgba(112,105,96,.24)", "width": 1.2}, hoverinfo="skip", showlegend=False))
    fig.add_trace(go.Scatter(
        x=[node["x"] for node in nodes],
        y=[node["y"] for node in nodes],
        text=[node["skill"] for node in nodes],
        mode="markers+text",
        textposition="top center",
        marker={
            "size": [min(34, 10 + node["count"] * 2) for node in nodes],
            "color": [color_by_type.get(node["skill_type"], "#706960") for node in nodes],
            "opacity": .92,
            "line": {"color": "#ffffff", "width": 1},
        },
        hovertemplate="%{text}<extra></extra>",
        showlegend=False,
    ))
    fig.update_layout(
        title={"text": "Semantic Association Clusters", "x": 0.02, "xanchor": "left"},
        margin={"l": 30, "r": 30, "t": 70, "b": 30},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"visible": False},
        yaxis={"visible": False, "scaleanchor": "x", "scaleratio": 1},
        height=620,
    )
    return fig


def methodology_pipeline_figure():
    go = plotly_go()
    if go is None:
        return None
    import math

    labels = ["Ingest", "Clean", "Extract", "Compare", "Validate", "Export"]
    details = [
        "courses, modules, jobs",
        "normalised evidence",
        "NER, aliases, confidence",
        "vectors, cosine, coverage",
        "scores, cells, equations",
        "DOCX, CSV, PNG assets",
    ]
    angles = [(-math.pi / 2) + index * (math.pi * 2 / len(labels)) for index in range(len(labels))]
    x_values = [math.cos(angle) for angle in angles]
    y_values = [math.sin(angle) for angle in angles]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=x_values + [x_values[0]],
        y=y_values + [y_values[0]],
        mode="lines",
        line={"color": "#f58220", "width": 5},
        hoverinfo="skip",
        showlegend=False,
    ))
    fig.add_trace(go.Scatter(
        x=x_values,
        y=y_values,
        mode="markers+text",
        marker={"size": 34, "color": ["#0b0b0b", "#706960", "#f58220", "#2f80ed", "#236b35", "#0b0b0b"]},
        text=[str(index + 1) for index in range(len(labels))],
        textfont={"color": "white", "size": 13},
        textposition="middle center",
        hoverinfo="skip",
        showlegend=False,
    ))
    for index, (label, detail) in enumerate(zip(labels, details)):
        fig.add_annotation(x=x_values[index] * 1.28, y=y_values[index] * 1.18, text=f"<b>{label}</b><br>{detail}", showarrow=False, align="center", font={"size": 13})
    fig.add_annotation(
        x=0,
        y=0,
        text="<b>Circular methodology loop</b><br>reviewed skills and aliases<br>feed the next refresh",
        showarrow=False,
        align="center",
        font={"size": 13},
        bgcolor="#fbfaf6",
        bordercolor="#e6ded3",
        borderpad=10,
    )
    fig.update_layout(
        title={"text": "Iterative Methodology Pipeline", "x": .02, "xanchor": "left"},
        width=1050,
        height=520,
        margin={"l": 20, "r": 20, "t": 60, "b": 30},
        plot_bgcolor="white",
        paper_bgcolor="white",
        xaxis={"visible": False, "range": [-1.7, 1.7]},
        yaxis={"visible": False, "range": [-1.4, 1.4], "scaleanchor": "x", "scaleratio": 1},
    )
    return fig


def crisp_dm_rows():
    return [
        ("Business Understanding", "Curriculum-job alignment objective", "Course-to-job gap question"),
        ("Data Understanding", "Courses, modules, jobs, sectors, dates", "Evidence audit"),
        ("Data Preparation", "Clean text, deduplicate, review skills and aliases", "Human-cleaned skill evidence"),
        ("Modelling", "NER, phrase matching, approved aliases, embeddings, ensemble scoring", "GapResult signals"),
        ("Evaluation", "Human Oversight approvals, validation equations, heatmaps", "Reviewed quality gate"),
        ("Deployment", "Dashboard visuals, CSV, PNG cache, DOCX export, hourly refresh", "Decision-support outputs"),
    ]


def crisp_dm_figure():
    go = plotly_go()
    if go is None:
        return None
    import math

    rows = crisp_dm_rows()
    angles = [(-math.pi / 2) + index * (math.pi * 2 / len(rows)) for index in range(len(rows))]
    x_values = [math.cos(angle) for angle in angles]
    y_values = [math.sin(angle) for angle in angles]
    colours = ["#0b0b0b", "#706960", "#f58220", "#2f80ed", "#236b35", "#7b4b16"]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=x_values + [x_values[0]],
        y=y_values + [y_values[0]],
        mode="lines",
        line={"color": "#d47722", "width": 4},
        hoverinfo="skip",
        showlegend=False,
    ))
    fig.add_trace(go.Scatter(
        x=x_values,
        y=y_values,
        mode="markers+text",
        marker={"size": 58, "color": colours, "line": {"color": "white", "width": 2}},
        text=[str(index + 1) for index in range(len(rows))],
        textfont={"color": "white", "size": 14},
        textposition="middle center",
        hovertext=[f"{row[0]}<br>{row[1]}<br>{row[2]}" for row in rows],
        hovertemplate="%{hovertext}<extra></extra>",
        showlegend=False,
    ))
    for index, row in enumerate(rows):
        fig.add_annotation(
            x=x_values[index] * 1.26,
            y=y_values[index] * 1.22,
            text=f"<b>{row[0]}</b><br>{row[2]}",
            showarrow=False,
            align="center",
            font={"size": 12},
        )
    fig.add_annotation(
        x=0,
        y=0,
        text="<b>Human-cleaned<br>data-mining loop</b><br>approved skills + aliases",
        showarrow=False,
        align="center",
        font={"size": 13},
        bgcolor="#fbfaf6",
        bordercolor="#e6ded3",
        borderpad=10,
    )
    fig.update_layout(
        title={"text": "CRISP-DM Data-Mining Process", "x": .02, "xanchor": "left"},
        width=900,
        height=620,
        margin={"l": 30, "r": 30, "t": 70, "b": 30},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"visible": False, "range": [-1.75, 1.75]},
        yaxis={"visible": False, "range": [-1.55, 1.55], "scaleanchor": "x", "scaleratio": 1},
    )
    return fig


def evidence_funnel_figure():
    go = plotly_go()
    if go is None:
        return None
    fig = go.Figure(go.Funnel(
        y=["Raw evidence", "Structured records", "Skill evidence", "Comparable signals", "Validation outputs"],
        x=[100, 82, 62, 42, 26],
        text=[
            "courses, jobs, scraped pages",
            "modules, adverts, dates, sectors",
            "entities, types, confidence",
            "embeddings, coverage, scores",
            "equations, cells, images",
        ],
        textposition="inside",
        marker={"color": ["#0b0b0b", "#3c3327", "#7b4b16", "#f58220", "#236b35"]},
        connector={"line": {"color": "#706960", "width": 1}},
        hoverinfo="skip",
    ))
    fig.update_layout(
        title={"text": "Evidence Funnel", "x": .02, "xanchor": "left"},
        width=850,
        height=430,
        margin={"l": 30, "r": 30, "t": 60, "b": 20},
        paper_bgcolor="white",
        font={"size": 13},
    )
    return fig


def validation_equation_figure():
    go = plotly_go()
    if go is None:
        return None
    labels = ["Semantic score", "Skill coverage", "Confidence", "Decision tree", "Pearson cells", "Divergence cells"]
    values = [75, 25, 15, 10, 100, 100]
    fig = go.Figure(go.Bar(
        x=labels,
        y=values,
        marker={"color": ["#0b0b0b", "#f58220", "#ffbd80", "#236b35", "#3f6fa9", "#c83a2f"]},
        text=["weighted", "weighted", "weighted", "weighted", "validated", "validated"],
        textposition="outside",
    ))
    fig.update_layout(
        title={"text": "Validation Calculations and Visual Cells", "x": .02, "xanchor": "left"},
        margin={"l": 65, "r": 30, "t": 70, "b": 120},
        plot_bgcolor="white",
        paper_bgcolor="white",
        yaxis={"title": "Relative validation role", "range": [0, 120], "gridcolor": "#f2eee8"},
        xaxis={"tickangle": -25},
        showlegend=False,
    )
    return fig


def transformer_architecture_rows():
    return [
        ("Raw text", "module/job sentence"),
        ("Subword tokenizer", "prefix/root/suffix pieces"),
        ("Token embeddings", "numeric token vectors"),
        ("Position encodings", "word order signal"),
        ("Transformer encoder", "self-attention + feed-forward"),
        ("Contextual states", "context-aware token vectors"),
        ("BIO classifier", "O, B-SKILL, I-SKILL"),
        ("Skill spans", "accepted entities + confidence"),
    ]


def runtime_event_flow_rows():
    return [
        ("User", "Task", "Queue analysis/export task"),
        ("Task", "Analysis", "Start background run"),
        ("Analysis", "DB", "Load courses, modules, jobs"),
        ("Analysis", "Scorer", "Prepare embeddings backend"),
        ("Analysis", "Extractor", "Extract module/job skills"),
        ("Extractor", "DB", "Persist skill_entities"),
        ("Analysis", "Ensemble", "Compute gaps and final scores"),
        ("Ensemble", "DB", "Save GapResult + SkillMatrix"),
        ("User", "Report", "Request paper export"),
        ("Report", "Visuals", "Render and cache PNG figures"),
        ("Report", "DOCX", "Embed figures, equations, tables"),
    ]


def school_summary_figure(visual_data):
    go = plotly_go()
    if go is None:
        return None

    rows = visual_data.get("school_summaries", [])
    labels = [row["school"] for row in rows]
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Average score", x=labels, y=[row["avg_score"] for row in rows], marker={"color": "#2f80ed"}))
    fig.add_trace(go.Bar(name="Missing evidence", x=labels, y=[row["missing_total"] for row in rows], marker={"color": "#f58220"}, yaxis="y2"))
    fig.update_layout(
        title={"text": "School Alignment Summary", "x": 0.02, "xanchor": "left"},
        margin={"l": 70, "r": 70, "t": 80, "b": 130},
        paper_bgcolor="white",
        plot_bgcolor="white",
        xaxis={"tickangle": -30},
        yaxis={"title": "Average score %", "gridcolor": "#f2eee8", "range": [0, 100]},
        yaxis2={"title": "Missing evidence", "overlaying": "y", "side": "right"},
        legend={"orientation": "h", "y": 1.08},
    )
    return fig


def oversight_counts():
    counts = Counter()
    for model in (JobAdvert, Module):
        for entities in model.objects.values_list("skill_entities", flat=True):
            for entity in entities or []:
                if isinstance(entity, dict):
                    counts[entity.get("label_status") or "machine"] += 1
    return counts


def build_research_paper_docx(run, visual_data):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    docx_add_field(footer, "PAGE", "1")

    results = list(GapResult.objects.filter(run=run).select_related("course", "job")) if run else []
    job_skill_rows = list(SkillMatrix.objects.filter(run=run, source="jobs").order_by("-frequency", "skill")[:25]) if run else []
    course_skill_rows = list(SkillMatrix.objects.filter(run=run, source="courses").order_by("-frequency", "skill")[:25]) if run else []
    data_export_rows = report_skill_entity_rows(run)
    schools = list(Course.objects.exclude(university_name="").values_list("university_name", flat=True).distinct())
    if Course.objects.filter(university_name="").exists():
        schools.append("Unassigned school")
    asset_dir = paper_artifact_dir(run)
    manifest_rows = []

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CurriculumMatch Research Paper Export")
    title_run.bold = True
    title_run.font.size = Pt(20)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Tables, graphs, visualisations, methodology notes, and human review evidence").italic = True
    document.add_paragraph(f"Generated: {timezone.localtime(timezone.now()).strftime('%d %B %Y %H:%M')}")
    document.add_paragraph(f"Analysis run: {run.name if run else 'No completed analysis run available'}")
    document.add_paragraph(f"Paper visual assets: {asset_dir.parent}")
    document.add_page_break()

    document.add_heading("Table of Contents", level=1)
    docx_add_field(document.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')
    document.add_paragraph("In Microsoft Word, right-click the table above and choose Update Field to refresh page numbers.")
    document.add_page_break()

    avg_score = round(sum(result.similarity_score for result in results) * 100 / len(results), 1) if results else 0
    counts = oversight_counts()
    alias_counts = Counter(SkillAlias.objects.values_list("status", flat=True))
    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "This enhanced export packages the system evidence into a research-paper style document. It includes the same dashboard "
        "source data as tables, plus server-side Plotly image exports for the main charts and visualisations."
    )
    add_key_value_table(document, [
        ("Schools", len(schools)),
        ("Courses", Course.objects.count()),
        ("Modules", Module.objects.count()),
        ("Job adverts", JobAdvert.objects.count()),
        ("Course-job comparisons", len(results)),
        ("Average final score", f"{avg_score}%"),
        ("Reviewed skill entities", counts.get("reviewed", 0)),
        ("Candidate skill entities", counts.get("candidate", 0)),
        ("Approved skill aliases", alias_counts.get("approved", 0)),
        ("Candidate skill aliases", alias_counts.get("candidate", 0)),
    ])

    document.add_heading("2. Methodology", level=1)
    document.add_paragraph(
        "A skill is treated as a named capability that can be evidenced in curriculum text or job adverts. Technical skills are "
        "tools, programming languages, platforms, methods, or measurable specialist techniques. Soft skills are human workplace "
        "capabilities such as communication, leadership, collaboration, and problem solving. Business or domain skills describe "
        "work-area knowledge such as accounting, compliance, marketing, risk, education, or healthcare."
    )
    document.add_paragraph(
        "The system combines rule-assisted extraction, dynamic learned skills, human review, and semantic comparison. A candidate "
        "skill can be suggested by the model, but it only becomes trusted learning data after a human approves or edits it."
    )
    document.add_paragraph(
        "The overall methodology is framed as a CRISP-DM data-mining process: business understanding defines the curriculum-to-market "
        "alignment question; data understanding inspects curriculum and job evidence; data preparation cleans and reviews skill "
        "entities and aliases; modelling applies NER, phrase matching, approved aliases, embeddings, and scoring; evaluation uses "
        "Human Oversight and validation visuals; deployment publishes dashboard visuals, CSV files, cached PNGs, and this paper export."
    )
    document.add_paragraph(
        "Human cleaning is therefore part of the data-mining loop rather than an afterthought. Reviewed skill entities and approved "
        "SkillAlias rows are treated as trusted data; the reviewed-alias refresh task can rerun analysis every hour so SkillMatrix "
        "frequencies, GapResult scores, dashboard graphics, downloadable visuals, and report figures are refreshed from the cleaned evidence."
    )
    document.add_paragraph(
        "The validation notebook and report builder make the scoring path auditable. Every stored GapResult can be inspected as "
        "a set of inputs, equations, and outputs: semantic similarity, explicit skill coverage, extraction confidence, decision-tree "
        "score, final ensemble score, correlation heatmap cells, role divergence bars, and skill gap cells."
    )
    document.add_paragraph(
        "The verified production flow is the Django background-task path implemented in analysis/services.py and analysis/tasks.py. "
        "The transformer architecture diagram documents the supported neural tagging pattern for contextual skill extraction, while "
        "the active pipeline still records the extractor backend used for each run so the report can distinguish learned NER, phrase "
        "matching, aliases, and fallback extraction."
    )
    add_simple_table(document, ["Order", "Implemented event", "Calculation or persisted output"], [
        ("1", "Create TaskRecord and AnalysisRun.", "Task status, run status, progress percent"),
        ("2", "Load courses, modules, and job adverts.", "module_map, job_map, analysis_text"),
        ("3", "Prepare SemanticSimilarityService.", "embeddings for module and job text"),
        ("4", "Ensure skill NER model and create SpacySkillExtractor.", "extractor backend plus skill_entities"),
        ("5", "Vectorise module and job evidence.", "stored vector, skills_extracted, skill_entities"),
        ("6", "Compare every eligible course-job pair.", "matched = C intersect J; missing = J - C; extra = C - J"),
        ("7", "Score semantic, coverage, confidence, and decision policy signals.", "s, k, c, t, final ensemble score"),
        ("8", "Persist GapResult and SkillMatrix rows.", "paper tables, heatmaps, role divergence, export visuals"),
        ("9", "Build research-paper export.", "cached PNG figures, manifest, DOCX equations, downloadable file"),
    ])
    add_simple_table(document, ["Step", "Calculation", "Output"], [
        ("Semantic match", "Mean of the top module-to-job cosine similarities", "semantic_score s"),
        ("Skill coverage", "|matchedSkills| / |jobSkills|", "skill_score k"),
        ("Matched confidence", "Mean of matched course/job entity confidence values", "confidence_score c"),
        ("Decision tree", "Transparent threshold rules over s, k, and c", "decision_tree_score t"),
        ("Final score", "(0.75s + 0.25k + 0.15c + 0.10t) / 1.25", "final_score"),
        ("Correlation cell", "Pearson r over binary role-profile skill presence vectors", "skill-skill heatmap cell"),
        ("Role divergence cell", "role skill share minus global skill baseline share", "diverging bar value"),
        ("Gap cell", "missing / (matched + missing), and matched / (matched + missing)", "gap % and coverage %"),
    ])

    document.add_heading("Methodology Diagrams", level=2)
    add_figure(
        document,
        "Methodology Diagram 0. CRISP-DM Data-Mining Process",
        "CRISP-DM frames the whole CurriculumMatch workflow from business question through human-cleaned extraction, modelling, evaluation, and deployed report assets.",
        crisp_dm_figure(),
        ["CRISP-DM phase", "CurriculumMatch implementation", "Audit output"],
        crisp_dm_rows(),
        asset_dir=asset_dir,
        image_name="methodology-crisp-dm-data-mining-process",
        manifest_rows=manifest_rows,
    )
    add_figure(
        document,
        "Methodology Diagram 1. End-to-End Iterative Pipeline",
        "Raw curriculum and job evidence is transformed into validated scores, visual cells, and paper assets; human review and deployment feed the next analysis cycle.",
        methodology_pipeline_figure(),
        ["Stage", "Action", "Output"],
        [
            ("1. Ingest", "Collect courses, modules, uploads, scraped pages, CSV/manual jobs, and Adzuna adverts.", "Raw evidence"),
            ("2. Clean", "Normalise text, parse job sections, handle stale tasks, and skip duplicates.", "Structured records"),
            ("3. Extract", "Apply spaCy NER, phrase matching, aliases, regex fallback, and phrase mining.", "Skill entities"),
            ("4. Compare", "Create embeddings and compute cosine similarity plus skill coverage.", "Signals"),
            ("5. Validate", "Expose equations, cells, score breakdowns, and notebook outputs.", "Audit evidence"),
            ("6. Export", "Render DOCX, CSV, PNG figures, and visual manifests.", "Paper assets"),
        ],
        asset_dir=asset_dir,
        image_name="methodology-pipeline",
        manifest_rows=manifest_rows,
    )
    add_figure(
        document,
        "Methodology Diagram 2. Evidence Funnel",
        "The pipeline narrows raw evidence into structured records, skill evidence, comparable signals, and validation outputs.",
        evidence_funnel_figure(),
        ["Layer", "Input", "Output"],
        [
            ("Raw evidence", "Course pages, module text, job adverts", "Input corpus"),
            ("Structured records", "Normalised text, dates, sectors", "Comparable records"),
            ("Skill evidence", "NER entities and reviewed labels", "Auditable skills"),
            ("Comparable signals", "Embeddings and coverage counts", "Scores"),
            ("Validation outputs", "Equations, cells, images, CSV", "Publication evidence"),
        ],
        asset_dir=asset_dir,
        image_name="methodology-evidence-funnel",
        manifest_rows=manifest_rows,
    )
    add_figure(
        document,
        "Methodology Diagram 3. Validation Calculations",
        "The validation layer exposes weighted model components and the visual cell calculations used in the paper.",
        validation_equation_figure(),
        ["Calculation", "Formula", "Output"],
        [
            ("Final score", "(0.75s + 0.25k + 0.15c + 0.10t) / 1.25", "GapResult score"),
            ("Coverage", "|matched| / |job skills|", "Skill score"),
            ("Correlation", "Pearson r over binary skill vectors", "Heatmap cell"),
            ("Divergence", "role share - baseline share", "Diverging bar"),
            ("Gap", "missing / (matched + missing)", "Gap percentage"),
        ],
        asset_dir=asset_dir,
        image_name="methodology-validation-calculations",
        manifest_rows=manifest_rows,
    )
    add_figure(
        document,
        "Methodology Diagram 4. Transformer Architecture for Skill Tagging",
        "The optional BERT/transformer layer converts text into contextual token states and predicts BIO skill labels.",
        None,
        ["Layer", "Role"],
        transformer_architecture_rows(),
        asset_dir=asset_dir,
        image_name="methodology-transformer-architecture",
        manifest_rows=manifest_rows,
    )
    add_figure(
        document,
        "Methodology Diagram 5. Runtime Flow of Events",
        "This sequence diagram follows the implemented background task, analysis, visual-cache, and DOCX export event order.",
        None,
        ["Source", "Target", "Event"],
        runtime_event_flow_rows(),
        asset_dir=asset_dir,
        image_name="methodology-runtime-flow-of-events",
        manifest_rows=manifest_rows,
    )

    document.add_heading("Supporting Equations", level=2)
    document.add_paragraph(
        "The equations below mirror the runtime implementation in analysis/semantic_similarity.py and analysis/services.py. "
        "They are written as Word equation objects plus plain-text descriptions so the syntax remains inspectable after export."
    )
    equations = [
        (
            "Cosine similarity",
            r"cos(a,b) = (a \cdot b) / (||a||_2 ||b||_2)",
            "Used by compute_similarity and then clamped into the [0, 1] scoring interval.",
        ),
        (
            "Course semantic score",
            r"s = mean(top_N({ cos(m_i, j) : m_i in course_modules }))",
            "The code sorts module-to-job similarities descending and averages the top TOP_MODULE_MATCH_COUNT values.",
        ),
        (
            "Skill coverage",
            r"k = |matchedSkills| / max(1, |uniqueJobSkills|)",
            "This is SemanticSimilarityService.skill_coverage_score(...).",
        ),
        (
            "Matched confidence",
            r"c = mean({ (conf_course(skill) + conf_job(skill)) / 2 : skill in matchedSkills })",
            "This is _matched_skill_confidence(...), using stored skill entity confidence values.",
        ),
        (
            "Decision policy fallback",
            r"t = max(0, min(1, 0.70s + 0.30k))",
            "Used when the interpretable threshold rules do not trigger a higher decision-tree score.",
        ),
        (
            "Final weighted score",
            r"final = (0.75s + 0.25k + 0.15c + 0.10t) / (0.75 + 0.25 + 0.15 + 0.10)",
            "Default weights come from SemanticSimilarityService unless overridden in settings.",
        ),
        (
            "Set gap calculation",
            r"matched = C \cap J; missing = J - C; extra = C - J",
            "This is the set algebra used by compute_gap(course_skills, job_skills).",
        ),
        (
            "Role divergence cell",
            r"d_{r,s} = share(skill_s in role_r) - share(skill_s in all_roles)",
            "Used by the role skill divergence visual to show emphasis above or below the job-market baseline.",
        ),
        (
            "Correlation heatmap cell",
            r"rho(x,y) = cov(x,y) / (sigma_x sigma_y)",
            "Used by the skill correlation heatmap over binary role-profile skill-presence vectors.",
        ),
        (
            "Transformer self-attention",
            r"Attention(Q,K,V) = softmax((QK^T) / sqrt(d_k))V",
            "Included to document the contextual token mechanism used by transformer-style skill tagging.",
        ),
    ]
    add_simple_table(document, ["Equation", "Code reference / interpretation"], [
        (name, description) for name, _equation, description in equations
    ])
    for name, equation, description in equations:
        document.add_paragraph(name).runs[0].bold = True
        docx_add_equation(document, equation)
        document.add_paragraph(description)

    document.add_heading("3. Visual Evidence", level=1)
    skill_line_rows = dashboard_skill_line_rows(run)
    if skill_line_rows:
        add_figure(
            document,
            "Dashboard Chart. Skill Demand vs Curriculum",
            "This is the paper export version of the Chart.js line chart on the executive dashboard.",
            dashboard_skill_line_figure(run),
            ["Skill", "Job evidence", "Course evidence"],
            skill_line_rows,
            asset_dir=asset_dir,
            image_name="dashboard-skill-demand-vs-curriculum",
            manifest_rows=manifest_rows,
        )
    if results:
        add_figure(
            document,
            "Dashboard Chart. Cosine Similarity Network",
            "This is the paper export version of the dashboard course-job similarity network.",
            similarity_network_figure(results),
            ["Course", "Job advert", "Score"],
            [
                (result.course.code or result.course.name, result.job.title, f"{result.similarity_percent}%")
                for result in sorted(results, key=lambda item: item.similarity_score, reverse=True)[:20]
            ],
            asset_dir=asset_dir,
            image_name="dashboard-cosine-similarity-network",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Dashboard Plotly Chart. Course-to-Job Similarity Cross-tab",
            "This is the paper export version of the dashboard Plotly cross-tab heatmap.",
            course_job_crosstab_figure(results),
            ["Course", "Job advert", "Score"],
            [
                (result.course.code or result.course.name, result.job.title, f"{result.similarity_percent}%")
                for result in sorted(results, key=lambda item: item.similarity_score, reverse=True)[:20]
            ],
            asset_dir=asset_dir,
            image_name="dashboard-course-job-crosstab",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Results Chart. Course Alignment Score-Band Heatmap",
            "This is the paper export version of the score-band heatmap shown on the analysis results page.",
            score_band_heatmap_figure(visual_data),
            ["Course", "Band", "Job count"],
            [
                (row["course"].code or row["course"].name, cell["band"], cell["count"])
                for row in visual_data.get("heatmap_rows", [])[:10]
                for cell in row.get("cells", [])
            ],
            asset_dir=asset_dir,
            image_name="results-course-alignment-score-band-heatmap",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Figure 1. Score Distribution",
            "This chart shows how course-to-job comparisons are distributed across score bands.",
            score_distribution_figure(results),
            ["Score band", "Comparisons"],
            score_distribution_rows(results),
            asset_dir=asset_dir,
            image_name="figure-01-score-distribution",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Figure 2. Skill Correlation Heatmap",
            "This heatmap shows Pearson correlation coefficients between skill presence vectors across role requirement profiles.",
            heatmap_figure(visual_data),
            ["Row skill", "Column skill", "Pearson r"],
            correlation_rows(visual_data),
            asset_dir=asset_dir,
            image_name="figure-02-skill-correlation-heatmap",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Figure 3. Matched vs Missing Skill Evidence",
            "Each point represents a course, with position showing matched and missing skill evidence.",
            scatter_figure(visual_data),
            ["Course", "Matched", "Missing", "Average score"],
            [(point["label"], point["x"], point["y"], point["score"]) for point in visual_data.get("scatter_points", [])],
            asset_dir=asset_dir,
            image_name="figure-03-matched-vs-missing",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Figure 4. School Summary",
            "This chart compares school-level average alignment with missing-skill evidence.",
            school_summary_figure(visual_data),
            ["School", "Average score", "Missing evidence"],
            [(row["school"], row["avg_score"], row["missing_total"]) for row in visual_data.get("school_summaries", [])],
            asset_dir=asset_dir,
            image_name="figure-04-school-summary",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Figure 5. Skill Gap Matrix",
            "This chart compares covered and missing evidence for high-demand skills.",
            skill_gap_figure(visual_data),
            ["Skill", "Demand", "Covered", "Missing", "Gap %"],
            skill_gap_rows(visual_data),
            asset_dir=asset_dir,
            image_name="figure-05-skill-gap-matrix",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Figure 6. Role Skill Divergence Overview",
            "This overview combines the first role profiles from the Plotly role divergence section.",
            role_divergence_figure(visual_data),
            ["Role", "Skill", "Divergence"],
            role_divergence_rows(visual_data),
            asset_dir=asset_dir,
            image_name="figure-06-role-skill-divergence",
            manifest_rows=manifest_rows,
        )
        for index, role in enumerate((visual_data.get("role_skill_divergence", {}).get("roles") or [])[:25], start=1):
            add_figure(
                document,
                f"Results Plotly Chart. Role Skill Divergence {index}: {role['role'][:60]}",
                "This is one of the individual role divergence Plotly bar charts from the analysis results page.",
                single_role_divergence_figure(role),
                ["Skill", "Divergence"],
                list(zip(role.get("skills", []), role.get("values", []))),
                asset_dir=asset_dir,
                image_name=f"results-role-skill-divergence-{index}-{role['role']}",
                manifest_rows=manifest_rows,
            )
    else:
        document.add_paragraph("No completed course-to-job analysis results are available for chart export yet.")

    if job_skill_rows:
        rows = top_skill_rows(job_skill_rows)
        labels, values = zip(*rows)
        add_figure(
            document,
            "Figure 7. Top Job Skills",
            "Top skills extracted from job-market evidence.",
            create_bar_figure("Top Job Skills", list(labels), list(values), "Skill", "Frequency", "#0b0b0b"),
            ["Skill", "Frequency"],
            rows,
            asset_dir=asset_dir,
            image_name="figure-07-top-job-skills",
            manifest_rows=manifest_rows,
        )
    if course_skill_rows:
        rows = top_skill_rows(course_skill_rows)
        labels, values = zip(*rows)
        add_figure(
            document,
            "Figure 8. Top Course Skills",
            "Top skills extracted from course and module evidence.",
            create_bar_figure("Top Course Skills", list(labels), list(values), "Skill", "Frequency", "#236b35"),
            ["Skill", "Frequency"],
            rows,
            asset_dir=asset_dir,
            image_name="figure-08-top-course-skills",
            manifest_rows=manifest_rows,
        )

    if data_export_rows:
        document.add_heading("Data Export Plotly Charts", level=2)
        add_figure(
            document,
            "Data Export Plotly Chart. Top Skill Evidence",
            "This is the paper export version of the Top Skill Evidence Plotly chart.",
            data_export_top_skills_figure(data_export_rows),
            ["Skill", "Records"],
            Counter(row["skill"] for row in data_export_rows).most_common(15),
            asset_dir=asset_dir,
            image_name="data-export-top-skill-evidence",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Data Export Plotly Chart. Skill Type and Source",
            "This is the paper export version of the Skill Type and Source Plotly chart.",
            data_export_type_source_figure(data_export_rows),
            ["Dimension", "Value", "Records"],
            [(kind, value, count) for kind, pairs in (
                ("Skill type", Counter(row["skill_type"] for row in data_export_rows).most_common(10)),
                ("Source", Counter(row["source_type"] for row in data_export_rows).most_common(10)),
            ) for value, count in pairs],
            asset_dir=asset_dir,
            image_name="data-export-skill-type-and-source",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Data Export Plotly Chart. NER Skill Evidence Density",
            "This is the paper export version of the source-by-type skill evidence heatmap.",
            data_export_skill_heatmap_figure(data_export_rows),
            ["Source", "Skill type", "Records"],
            [(source, skill_type, count) for (source, skill_type), count in Counter((row["source_type"], row["skill_type"]) for row in data_export_rows).items()],
            asset_dir=asset_dir,
            image_name="data-export-ner-skill-evidence-density",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Data Export Plotly Chart. Skills Linked to Type and Source",
            "This is the paper export version of the skill evidence network on the Data Export page.",
            data_export_skill_network_figure(data_export_rows),
            ["Skill", "Type", "Source"],
            [(row["skill"], row["skill_type"], row["source_type"]) for row in data_export_rows[:30]],
            asset_dir=asset_dir,
            image_name="data-export-skill-network",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Data Export Plotly Chart. Skill Evidence Trend and One-Year Forecast",
            "This is the paper export version of the skill demand forecast chart. When no dated evidence exists, the source rows are shown instead.",
            data_export_forecast_figure(data_export_rows),
            ["Skill", "Year", "Records"],
            [(skill, year, count) for (skill, year), count in Counter((row["skill"], row.get("extracted_year")) for row in data_export_rows if row.get("extracted_year")).most_common(30)],
            asset_dir=asset_dir,
            image_name="data-export-skill-forecast",
            manifest_rows=manifest_rows,
        )
        add_figure(
            document,
            "Data Export Plotly Chart. Semantic Association Clusters",
            "This is the paper export version of the semantic skill cluster chart.",
            data_export_cluster_figure(data_export_rows),
            ["Skill", "Records"],
            Counter(row["skill"] for row in data_export_rows).most_common(30),
            asset_dir=asset_dir,
            image_name="data-export-semantic-association-clusters",
            manifest_rows=manifest_rows,
        )

    document.add_heading("4. Tables", level=1)
    document.add_heading("School Summary Table", level=2)
    add_simple_table(
        document,
        ["School", "Average score", "Matched evidence", "Missing evidence", "Courses", "Risk"],
        [
            [
                row["school"],
                f'{row["avg_score"]}%',
                row["matched_total"],
                row["missing_total"],
                row["course_count"],
                "Yes" if row["mismatch"] else "No",
            ]
            for row in visual_data.get("school_summaries", [])
        ] or [["No analysis data", "", "", "", "", ""]],
    )

    document.add_heading("Skill Gap Matrix Table", level=2)
    add_simple_table(
        document,
        ["Skill", "Demand evidence", "Covered", "Missing", "Gap %"],
        skill_gap_rows(visual_data, limit=30) or [["No matrix data", "", "", "", ""]],
    )

    document.add_heading("Top Course-to-Job Results", level=2)
    top_results = sorted(results, key=lambda item: item.similarity_score, reverse=True)[:25]
    add_simple_table(
        document,
        ["Course", "Job advert", "Company", "Score", "Matched skills", "Missing skills"],
        [
            [
                result.course.name[:60],
                result.job.title[:60],
                result.job.company or "",
                f"{result.similarity_percent}%",
                ", ".join((result.matched_skills or [])[:6]),
                ", ".join((result.missing_skills or [])[:6]),
            ]
            for result in top_results
        ] or [["No analysis data", "", "", "", "", ""]],
    )

    document.add_heading("5. Human Oversight and Learning", level=1)
    document.add_paragraph(
        "The human oversight workflow keeps model suggestions auditable. Reviewers can approve, edit, reject, or delete candidate "
        "skills. Approved rows update the trusted skill lists and can be included in the next NER training run."
    )
    add_key_value_table(document, [
        ("Candidate entities waiting", counts.get("candidate", 0)),
        ("Reviewed entities", counts.get("reviewed", 0)),
        ("Excluded entities", counts.get("exclude", 0)),
        ("Machine/legacy entities", counts.get("machine", 0) + counts.get("legacy", 0)),
    ])

    document.add_heading("6. Limitations", level=1)
    document.add_paragraph(
        "The charts are evidence summaries, not automatic curriculum decisions. Scores can be affected by short module descriptions, "
        "thin job adverts, noisy text, duplicate wording, sector bias, or terminology differences. Human review remains necessary "
        "before changing curriculum content or retraining the model."
    )

    output = BytesIO()
    write_visual_manifest(asset_dir, manifest_rows)
    document.save(output)
    output.seek(0)
    return output
